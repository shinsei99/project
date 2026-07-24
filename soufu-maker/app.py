import streamlit as st
import subprocess
import json
import os
import io
import shutil
import html as html_mod
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _find_claude() -> str:
    """環境に依らず claude 実行ファイルを探す（PATH → よくある場所の順）。"""
    found = shutil.which("claude")
    if found:
        return found
    for p in (
        os.path.expanduser("~/.local/bin/claude"),
        "/opt/homebrew/bin/claude",
        "/usr/local/bin/claude",
    ):
        if os.path.exists(p):
            return p
    return "claude"  # 最後の手段（PATH に任せる）


CLAUDE_BIN = _find_claude()
CLAUDE_TIMEOUT = 120
TANTOU_FILE  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tantou.json")

# 会社は「大京商事株式会社」で固定（住所・TEL・FAX は送付書ジェネレーターの大京情報を継承）
FIXED_COMPANY = {
    "company": "大京商事株式会社",
    "address": "大阪市都島区東野田町2-3-14",
    "tel": "０６－６３５３－０４１８",
    "fax": "０６－６３５３－０２８０",
}

# 担当者マスタ（切り替え対象）。担当情報は「名前」と「メールアドレス」のみ。
DEFAULT_TANTOU = [
    {"name": "鷲見　慎一", "email": ""},
]

PREAMBLE = (
    "拝啓　時下いよいよご清祥のこととお慶び申し上げます。\n"
    "平素は何かとご高配を賜り有難く厚くお礼申し上げます。\n"
    "さて、掲題につき下記の通り添付申し上げますので\n"
    "ご査収下さいますようお願い申し上げます。　　　　敬具"
)

_FW  = str.maketrans("0123456789", "０１２３４５６７８９")
# 全角数字・記号 → 半角（TEL/FAX 表示用）
_HALF = str.maketrans("０１２３４５６７８９－（）　", "0123456789-() ")

_ASCII_FONT = "Times New Roman"   # 英数字フォント
_JP_FONT    = "ＭＳ 明朝"           # 日本語フォント（Windows標準の明朝）


def to_halfwidth(s: str) -> str:
    return (s or "").translate(_HALF)


# ── Persistence ───────────────────────────────────────────────────────────────
def load_tantou() -> list:
    if os.path.exists(TANTOU_FILE):
        try:
            with open(TANTOU_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return [t.copy() for t in DEFAULT_TANTOU]


def save_tantou(tantou: list):
    with open(TANTOU_FILE, "w", encoding="utf-8") as f:
        json.dump(tantou, f, ensure_ascii=False, indent=2)


def build_sender(tantou: dict) -> dict:
    """固定会社情報 + 選択中の担当（名前・メール）から送付者 dict を組み立てる。"""
    return {
        "company": FIXED_COMPANY["company"],
        "title": "",
        "name": tantou.get("name", ""),
        "address": FIXED_COMPANY["address"],
        "tel": FIXED_COMPANY["tel"],
        "fax": FIXED_COMPANY["fax"],
        "email": tantou.get("email", ""),
    }


# ── Date ──────────────────────────────────────────────────────────────────────
def to_reiwa(d: date) -> str:
    y = d.year - 2018
    return f"令和{str(y).translate(_FW)}年{str(d.month).translate(_FW)}月{str(d.day).translate(_FW)}日"


# ── AI Generation ─────────────────────────────────────────────────────────────
def generate_body(recipient: str, memo: str, sender: dict) -> str:
    sender_name = f"{sender.get('title', '')}　{sender.get('name', '')}".strip()
    sender_desc = f"{sender['company']} {sender_name}" if sender.get("company") else sender_name

    prompt = f"""ビジネス文書の専門家として、送付状の「記」以下に書く本文メッセージを作成してください。

■ 送付先: {recipient}
■ 送付者: {sender_desc}
■ 送付の概要・用件: {memo.strip() if memo.strip() else "（詳細未入力）"}

【作成ルール】
- 「お世話になります。」で書き出す
- 何を送付するか、理由・背景を自然な流れで1〜3文にまとめる
- 締めは「よろしくお願いいたします。」など状況に合わせて
- 余計な前置きや署名は一切不要
- 送付状らしい簡潔・丁寧な文体
- 本文のみを出力すること

本文のみを出力してください:"""

    cmd = [CLAUDE_BIN, "-p", prompt, "--output-format", "json",
           "--dangerously-skip-permissions", "--model", "sonnet"]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=CLAUDE_TIMEOUT)
    except FileNotFoundError:
        raise RuntimeError("`claude` コマンドが見つかりません。")
    except subprocess.TimeoutExpired:
        raise RuntimeError(f"生成が{CLAUDE_TIMEOUT}秒を超えたため中断しました。")
    if proc.returncode != 0:
        raise RuntimeError(f"claude コマンドが失敗しました\n{proc.stderr.strip()[:300]}")
    result = json.loads(proc.stdout)
    if result.get("is_error"):
        raise RuntimeError(f"Claude がエラーを返しました: {result.get('result')}")
    return result.get("result", "").strip()


# ── Word Document（python-docx でゼロから生成 = 常に正当な OOXML）──────────────
# 旧実装は変換ツール由来の非標準テンプレート(template.docx)を差し込む方式で、
# 不正な要素/属性名により Word が「破損 → 開いて修復」を要求していた。
# python-docx で新規生成すれば styles.xml 等を含む正当なパッケージになり破損しない。
# レイアウトは完成イメージプレビュー（render_preview）に合わせている。

def _sender_lines(sender: dict) -> list:
    """プレビューと同一の送付者行 [(テキスト, 太字), ...]（TEL/FAXは半角）。"""
    lines = []
    if sender.get("company"):
        lines.append((sender["company"], True))
    tn = "　".join(filter(None, [sender.get("title", ""), sender.get("name", "")]))
    if tn:
        lines.append((tn, False))
    if sender.get("address"):
        lines.append((sender["address"], False))
    if sender.get("tel"):
        lines.append((f"TEL  {to_halfwidth(sender['tel'])}", False))
    if sender.get("fax"):
        lines.append((f"FAX  {to_halfwidth(sender['fax'])}", False))
    if sender.get("email"):
        lines.append((f"MAIL  {sender['email']}", False))
    return lines


def _style_run(run, size_pt: float, bold: bool = False):
    run.font.size = Pt(size_pt)
    run.font.name = _ASCII_FONT
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), _JP_FONT)


def _add_para(doc, text: str, size_pt: float, align, bold: bool = False,
              space_after: float = 4, space_before: float = 0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = 1.15
    _style_run(p.add_run(text), size_pt, bold)
    return p


def _add_rule_borders(p):
    """段落の上下に罫線（「書類送付の件」の囲み枠を再現）。"""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "4")
        e.set(qn("w:color"), "555555")
        pbdr.append(e)
    pPr.append(pbdr)


def create_docx(recipient: str, doc_date: date, body: str, sender: dict) -> io.BytesIO:
    """完成イメージプレビュー準拠の送付状を python-docx で新規生成して返す。"""
    R = WD_ALIGN_PARAGRAPH.RIGHT
    L = WD_ALIGN_PARAGRAPH.LEFT
    C = WD_ALIGN_PARAGRAPH.CENTER

    doc = Document()

    # ページ設定（Letter 8.5×11 / 余白 上下1" 左右1.25"）
    sec = doc.sections[0]
    sec.page_width   = Inches(8.5)
    sec.page_height  = Inches(11)
    sec.top_margin   = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin  = Inches(1.25)
    sec.right_margin = Inches(1.25)

    # 既定フォント（Normal スタイル）にも日本語フォントを設定
    normal = doc.styles["Normal"]
    normal.font.name = _ASCII_FONT
    normal.font.size = Pt(12)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _JP_FONT)

    _add_para(doc, to_reiwa(doc_date), 12, R)                       # 日付（右）
    _add_para(doc, f"{recipient}　様", 16, L, bold=True, space_after=10)  # 宛名（左・大）

    for text, bold in _sender_lines(sender):                        # 送付者（右寄せ）
        _add_para(doc, text, 12, R, bold=bold, space_after=2)

    _add_para(doc, "", 12, L, space_after=8)                        # 余白

    # 「書類送付の件」を全角スペースで等間隔に区切り中央寄せ
    title = _add_para(doc, "　".join("書類送付の件"), 14, C, bold=True, space_after=12)
    _add_rule_borders(title)                                        # 上下罫線

    for line in PREAMBLE.split("\n"):                               # 前文（中央）
        _add_para(doc, line, 12, C, space_after=2)

    # 敬具 → 記 → 本文 の間に行間を空ける
    _add_para(doc, "記", 13, C, bold=True, space_before=18, space_after=16)

    for line in [l for l in body.split("\n") if l.strip()]:         # 本文（左・やや大きめ）
        _add_para(doc, line, 13, L, space_after=6)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ── Preview HTML ──────────────────────────────────────────────────────────────
def render_preview(recipient: str, doc_date: date, body: str, sender: dict):
    e = html_mod.escape

    s_lines = []
    if sender.get("company"):
        s_lines.append(f"<b>{e(sender['company'])}</b>")
    tn = "　".join(filter(None, [sender.get("title", ""), sender.get("name", "")]))
    if tn:
        s_lines.append(e(tn))
    if sender.get("address"):
        s_lines.append(e(sender["address"]))
    if sender.get("tel"):
        s_lines.append(f"TEL&nbsp;&nbsp;{e(to_halfwidth(sender['tel']))}")
    if sender.get("fax"):
        s_lines.append(f"FAX&nbsp;&nbsp;{e(to_halfwidth(sender['fax']))}")
    if sender.get("email"):
        s_lines.append(f"MAIL&nbsp;&nbsp;{e(sender['email'])}")

    st.markdown(
        f"""
<div style="
    background:white;border:1px solid #ccc;border-radius:6px;
    font-family:'Hiragino Mincho Pro','MS 明朝','Yu Mincho',serif;
    color:#111;box-shadow:0 3px 12px rgba(0,0,0,0.1);
    max-width:680px;margin:0 auto;padding:28px 36px 24px 36px;
    line-height:2;font-size:13px;">
  <div style="text-align:right;margin-bottom:10px;">{e(to_reiwa(doc_date))}</div>
  <div style="margin-bottom:6px;font-size:17px;font-weight:bold;">{e(recipient)}　様</div>
  <div style="margin-bottom:16px;line-height:1.9;text-align:right;">{"<br>".join(s_lines)}</div>
  <div style="text-align:center;font-size:15px;font-weight:bold;letter-spacing:2px;
      border-top:1px solid #555;border-bottom:1px solid #555;
      padding:6px 0;margin-bottom:14px;">{"　".join("書類送付の件")}</div>
  <div style="margin-bottom:14px;text-align:center;">{e(PREAMBLE).replace(chr(10),"<br>")}</div>
  <div style="text-align:center;font-size:14px;font-weight:bold;margin:14px 0 10px;">記</div>
  <div style="margin-bottom:20px;">{e(body).replace(chr(10),"<br>")}</div>
</div>
""",
        unsafe_allow_html=True,
    )


# ── Sidebar: 担当 Management ───────────────────────────────────────────────────
def sidebar_tantou() -> dict:
    tantou = st.session_state.setdefault("tantou", load_tantou())

    # 会社は固定表示（切り替え不可）
    st.markdown("## 🏢 送付元（固定）")
    st.info(
        f"**{FIXED_COMPANY['company']}**\n\n"
        f"{FIXED_COMPANY['address']}\n\n"
        f"TEL {to_halfwidth(FIXED_COMPANY['tel'])}　/　FAX {to_halfwidth(FIXED_COMPANY['fax'])}"
    )

    st.markdown("## 👤 担当を選択")
    names = [t.get("name", "") or "（無名）" for t in tantou]
    idx = st.session_state.get("tantou_idx", 0)
    if idx >= len(tantou):
        idx = 0

    chosen_name = st.radio("担当", names, index=idx, key="tantou_radio",
                           label_visibility="collapsed")
    chosen_idx = names.index(chosen_name)
    st.session_state["tantou_idx"] = chosen_idx
    cur = tantou[chosen_idx]
    if cur.get("email"):
        st.caption(f"✉️ {cur['email']}")

    st.divider()
    with st.expander("✏️ 担当を編集・追加", expanded=False):
        mode = st.radio("操作", ["選択中を編集", "新規追加"], horizontal=True, key="edit_mode")
        et = tantou[chosen_idx].copy() if mode == "選択中を編集" else {"name": "", "email": ""}

        # 新規追加時は白紙、編集時は担当ごとに、widget key を切り替えて value を確実に反映
        suffix = "new" if mode == "新規追加" else f"edit_{chosen_idx}"
        e_name  = st.text_input("名前",         value=et.get("name", ""),  key=f"e_name_{suffix}")
        e_email = st.text_input("メールアドレス", value=et.get("email", ""), key=f"e_email_{suffix}")

        new_data = {"name": e_name.strip(), "email": e_email.strip()}

        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 保存", use_container_width=True):
                if not e_name.strip():
                    st.warning("名前を入力してください")
                else:
                    if mode == "新規追加":
                        tantou.append(new_data)
                        st.session_state["tantou_idx"] = len(tantou) - 1
                    else:
                        tantou[chosen_idx] = new_data
                    save_tantou(tantou)
                    st.session_state["tantou"] = tantou
                    st.success("保存しました")
                    st.rerun()
        with col2:
            if mode == "選択中を編集" and len(tantou) > 1:
                if st.button("🗑 削除", use_container_width=True, type="secondary"):
                    tantou.pop(chosen_idx)
                    st.session_state["tantou_idx"] = 0
                    save_tantou(tantou)
                    st.session_state["tantou"] = tantou
                    st.rerun()

    return build_sender(tantou[chosen_idx])


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    st.set_page_config(page_title="送付書メーカー", page_icon="📮", layout="wide")

    with st.sidebar:
        sender = sidebar_tantou()

    st.title("📮 送付書メーカー")
    st.caption("宛名と用件を入力するだけで、AIが送付状を自動生成 → Word（.docx）ダウンロード")
    st.divider()

    left, right = st.columns([1, 1.2], gap="large")

    with left:
        st.subheader("📝 送付内容を入力")
        recipient = st.text_input("① 宛名", placeholder="例：山田　太郎",
                                  help="「様」は自動で付きます。")
        doc_date  = st.date_input("② 日付", value=date.today())
        memo = st.text_area(
            "③ 送付内容・用件メモ",
            placeholder="例：\n・契約書類一式をお送りします\n・先日お話しした見積書を添付します",
            height=150,
            help="箇条書きでもOK。AIが自然な文章に仕上げます。",
        )
        gen_btn = st.button("✨ 送付書を生成する", type="primary", use_container_width=True)

    with right:
        st.subheader("📄 プレビュー・ダウンロード")

        if gen_btn:
            if not recipient.strip():
                st.warning("⚠️ 宛名を入力してください。")
            elif not memo.strip():
                st.warning("⚠️ 送付内容・用件メモを入力してください。")
            else:
                with st.spinner("AIが本文を生成しています..."):
                    try:
                        body = generate_body(recipient.strip(), memo, sender)
                        st.session_state["body"]        = body
                        st.session_state["meta"]        = {"recipient": recipient.strip(),
                                                           "doc_date": doc_date}
                        st.session_state["has_content"] = True
                        st.success("✅ 生成完了！")
                    except Exception as ex:
                        st.error(f"❌ エラーが発生しました: {ex}")

        if st.session_state.get("has_content"):
            meta   = st.session_state["meta"]
            edited = st.text_area("✏️ 本文（記 以下）を確認・編集", height=160, key="body",
                                  help="自由に編集できます。")

            st.markdown("---")
            with st.expander("📋 完成イメージ（プレビュー）", expanded=True):
                render_preview(meta["recipient"], meta["doc_date"], edited, sender)

            st.markdown("---")
            try:
                buf = create_docx(meta["recipient"], meta["doc_date"], edited, sender)
                from datetime import datetime
                fname = f"送付書_{meta['recipient']}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.download_button(
                    label="📥 Word（.docx）をダウンロード",
                    data=buf,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary",
                )
                st.caption(fname)
            except Exception as ex:
                st.error(f"Word生成エラー: {ex}")

        elif not gen_btn:
            st.info("👈 左側で宛名・用件を入力し、「送付書を生成する」ボタンを押してください。")


if __name__ == "__main__":
    main()
