# -*- coding: utf-8 -*-
"""買取DMメーカー — 所有者台帳（空地調査）Excel から未活用地・空き家の買取DMを差し込み量産。

・平野区空地調査フォーマット（登記名義人／現住所／所在・地番／地目／地積…）を想定
・列は自動マッピング（見出し名で推定）、手動で上書き可
・現住所なし除外・宛先重複集約・個人/法人フィルタ・抵当権フラグ除外
・1通1ページの結合docx、または名義人ごとの個別docx(ZIP) を出力
"""
import io
import os
import re
import sys
import json
import shutil
import zipfile
from datetime import date

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SENDERS_FILE = os.path.join(HERE, "senders.json")

# 謄本読み取りは既存アプリ（baikai-generator）のパーサを再利用する
BAIKAI_SERVICES = os.path.join(os.path.dirname(HERE), "baikai-generator", "services")

# 台帳の確定フォーマット（1物件=1行・土地列＋建物列）
CANON_COLS = ["NO", "市", "所在", "地番", "地目", "地積・㎡",
              "建物種類", "建物構造", "床面積・㎡",
              "登記名義人", "持分", "郵便番号", "現住所", "電話番号", "備考"]

# 文面（固定）
HEADLINE = "空き家・未活用地を、諸経費ゼロで買い取ります。"
SUBJECT_KIND = "未活用地や空き家"

_ASCII = "Times New Roman"
_MIN = "ＭＳ 明朝"
_GOT = "ＭＳ ゴシック"
R = WD_ALIGN_PARAGRAPH.RIGHT
L = WD_ALIGN_PARAGRAPH.LEFT
C = WD_ALIGN_PARAGRAPH.CENTER
JU = WD_ALIGN_PARAGRAPH.JUSTIFY

CORP_MARKS = ("㈱", "株式会社", "有限会社", "（株）", "(株)", "㈲", "合同会社", "財団", "社団")

# 台帳の標準列名 → 内部キー（別名も許容）
COLMAP = {
    "name":   ["登記名義人", "名義人", "所有者", "氏名"],
    "addr":   ["現住所", "送付先住所", "住所"],
    "postal": ["郵便番号", "〒番号", "郵便"],
    "city":   ["市", "市区町村"],
    "basho":  ["所在", "所在地", "町名"],
    "chiban": ["地番"],
    "chimoku":["地目"],
    "area":   ["地積・㎡", "地積", "地積(㎡)", "面積"],
    "kyoyu":  ["共有"],
    "biko":   ["備考", "摘要"],
}


# ─────────────────────────── データ ───────────────────────────
# 既定の差出人（senders.json が無い場合のフォールバック）。
# 編集後は senders.json に保存され、そちらが優先される（senders.json は .gitignore 済み）。
DEFAULT_SENDERS = [
    {
        "label": "新誠プロパティマネジメント（買取主体）",
        "company": "新誠プロパティマネジメント株式会社",
        "title": "代表取締役", "name": "鷲見　慎一",
        "address": "大阪市北区大淀中3-1-15", "postal": "531-0076",
        "tel": "０６－６９３５－７２６７", "fax": "０６－７６３５－７８１１",
        "mobile": "090-8530-0184", "email": "info@shinsei-pm.co.jp",
    },
    {
        "label": "大京商事",
        "company": "大京商事株式会社",
        "title": "専務取締役", "name": "鷲見　慎一",
        "address": "大阪市都島区東野田町2-3-14", "postal": "",
        "tel": "０６－６３５３－０４１８", "fax": "０６－６３５３－０２８０",
        "mobile": "", "email": "shin@daikyocorp.co.jp",
    },
]


def load_senders():
    if os.path.exists(SENDERS_FILE):
        try:
            with open(SENDERS_FILE, encoding="utf-8") as f:
                data = json.load(f)
            if data:
                return data
        except Exception:
            pass
    return json.loads(json.dumps(DEFAULT_SENDERS))


def save_senders(senders):
    with open(SENDERS_FILE, "w", encoding="utf-8") as f:
        json.dump(senders, f, ensure_ascii=False, indent=2)


SENDER_FIELDS = [
    ("label", "表示名（ラベル）"), ("company", "会社名"), ("title", "役職"),
    ("name", "氏名"), ("postal", "郵便番号"), ("address", "住所"),
    ("tel", "TEL"), ("fax", "FAX"), ("mobile", "携帯"), ("email", "メール"),
]


# ─────────────────── 謄本 → 台帳 ───────────────────
def _get_registry_parser():
    """baikai-generator の registry_parser を読み込む（claudeパスは実機に合わせる）。"""
    if BAIKAI_SERVICES not in sys.path:
        sys.path.insert(0, BAIKAI_SERVICES)
    import registry_parser as rp
    claude = shutil.which("claude") or rp.CLAUDE_BIN
    rp.CLAUDE_BIN = claude
    return rp


def _num(s):
    if not s:
        return ""
    m = re.search(r"[\d,]+(?:\.\d+)?", str(s))
    return float(m.group().replace(",", "")) if m else ""


def split_city(s):
    """所在文字列から市区町村を切り出す。('加東市秋津字西山') → ('加東市','秋津字西山')。"""
    s = (s or "").strip()
    if not s:
        return "", ""
    s = re.sub(r"^.{2,3}[都道府県]", "", s)          # 先頭の都道府県を除去
    m = re.match(r"^(.+?市.+?区|.+?郡.+?[町村]|.+?[市区町村])", s)
    if m:
        return m.group(1), s[m.end():]
    return "", s


def registry_to_record(parsed, next_no=1):
    """registry_parser の出力辞書 → 台帳1行（CANON_COLS準拠）。"""
    tochi = parsed.get("土地", {}) or {}
    tate = parsed.get("建物", {}) or {}
    name = parsed.get("所有者氏名") or parsed.get("登記名義人氏名") or ""
    addr = parsed.get("所有者住所") or parsed.get("登記名義人住所") or ""
    raw_shozai = tochi.get("所在") or parsed.get("物件所在地") or tate.get("所在") or ""
    city, shozai = split_city(raw_shozai)
    biko = []
    if parsed.get("抵当権"):
        biko.append(f"抵当権:{parsed['抵当権']}")
    if tochi.get("権利"):
        biko.append(f"権利:{tochi['権利']}")
    return {
        "NO": next_no,
        "市": city,
        "所在": shozai,
        "地番": tochi.get("地番", ""),
        "地目": tochi.get("地目", ""),
        "地積・㎡": _num(tochi.get("地積", "")),
        "建物種類": tate.get("種類", ""),
        "建物構造": tate.get("構造", ""),
        "床面積・㎡": _num(tate.get("延床面積") or tate.get("床面積", "")),
        "登記名義人": name,
        "持分": "1/1" if name else "",
        "郵便番号": "",
        "現住所": addr,
        "電話番号": "",
        "備考": " ／ ".join(biko),
    }


def daicho_to_xlsx(df):
    """台帳DataFrame → 体裁付きxlsx（土地/建物を色分け）。"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    cols = [c for c in CANON_COLS if c in df.columns] + \
           [c for c in df.columns if c not in CANON_COLS]
    wb = Workbook()
    ws = wb.active
    ws.title = "NO順"
    navy = "FF1F3A5F"
    land = "FFEAF1F8"
    bldg = "FFF3EEE6"
    thin = Side(style="thin", color="FFBBBBBB")
    box = Border(left=thin, right=thin, top=thin, bottom=thin)
    got = "ＭＳ ゴシック"
    land_h = {"地目", "地積・㎡"}
    bldg_h = {"建物種類", "建物構造", "床面積・㎡"}
    left_h = {"市", "所在", "登記名義人", "現住所", "備考"}
    for j, h in enumerate(cols, 1):
        c = ws.cell(1, j, h)
        c.font = Font(name=got, size=9, bold=True, color="FFFFFFFF")
        c.fill = PatternFill("solid", fgColor=navy)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = box
    ws.row_dimensions[1].height = 28
    for i, (_, row) in enumerate(df.iterrows()):
        r = 2 + i
        for j, h in enumerate(cols, 1):
            val = row.get(h, "")
            if pd.isna(val):
                val = ""
            c = ws.cell(r, j, val)
            c.border = box
            c.font = Font(name=got, size=9)
            c.alignment = Alignment(horizontal=("left" if h in left_h else "center"),
                                    vertical="center", wrap_text=(h == "備考"))
            if h in ("地積・㎡", "床面積・㎡") and isinstance(val, (int, float)):
                c.number_format = "#,##0.00"
            if h in land_h:
                c.fill = PatternFill("solid", fgColor=land)
            elif h in bldg_h:
                c.fill = PatternFill("solid", fgColor=bldg)
    wmap = {"NO": 4, "市": 7, "所在": 12, "地番": 12, "地目": 8, "地積・㎡": 8,
            "建物種類": 9, "建物構造": 20, "床面積・㎡": 9, "登記名義人": 22,
            "持分": 6, "郵便番号": 8, "現住所": 20, "電話番号": 12, "備考": 40}
    for j, h in enumerate(cols, 1):
        ws.column_dimensions[get_column_letter(j)].width = wmap.get(h, 12)
    ws.freeze_panes = "A2"
    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


def to_halfwidth(s):
    return (s or "").translate(str.maketrans("０１２３４５６７８９－（）　", "0123456789-() "))


def clean(v):
    if v is None:
        return ""
    if isinstance(v, float):
        if v != v:  # NaN
            return ""
        if v == int(v):
            v = int(v)
    s = str(v).strip()
    return "" if s.lower() == "nan" else s


def is_corp(name):
    return any(m in name for m in CORP_MARKS)


def auto_map(columns):
    """DataFrameの列名から内部キー→実列名 を推定。
    完全一致を最優先し、無ければ部分一致（別名の登場順を優先）。"""
    cols = [str(c).strip() for c in columns]
    mapping = {}
    # 1) 完全一致
    for key, aliases in COLMAP.items():
        for a in aliases:
            if a in cols:
                mapping[key] = columns[cols.index(a)]
                break
    # 2) 部分一致（未確定のキーのみ）
    used = set(mapping.values())
    for key, aliases in COLMAP.items():
        if key in mapping:
            continue
        for a in aliases:
            hit = next((columns[i] for i, cc in enumerate(cols)
                        if a in cc and columns[i] not in used), None)
            if hit is not None:
                mapping[key] = hit
                used.add(hit)
                break
    return mapping


def build_records(df, m, require_addr=True):
    recs = []
    for _, row in df.iterrows():
        def g(key):
            col = m.get(key)
            return clean(row[col]) if col and col in df.columns else ""
        name = g("name")
        addr = g("addr")
        if not name:
            continue
        if require_addr and (addr in ("", "該当なし", "不明")):
            continue
        recs.append({
            "name": name, "addr": addr, "postal": g("postal"),
            "city": g("city"), "basho": g("basho"), "chiban": g("chiban"),
            "chimoku": g("chimoku"), "area": g("area"),
            "kyoyu": g("kyoyu"), "biko": g("biko"),
        })
    return recs


def dedupe(recs):
    seen, uniq = {}, []
    for r in recs:
        key = (r["name"], r["addr"])
        if key in seen:
            seen[key]["extra"] += 1
            continue
        r["extra"] = 0
        seen[key] = r
        uniq.append(r)
    return uniq


# ─────────────────────────── docx ───────────────────────────
def to_reiwa(d):
    y = d.year - 2018
    return f"令和{'元' if y == 1 else y}年{d.month}月{d.day}日"


NAVY = RGBColor(0x1F, 0x3A, 0x5F)


def srun(run, size, bold=False, jp=_MIN, color=None):
    run.font.size = Pt(size)
    run.font.name = _ASCII
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), jp)


def para(doc, text, size, align, bold=False, after=4, before=0, ls=1.3,
         indent=False, jp=_MIN, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = ls
    if indent:
        pf.first_line_indent = Pt(size)
    srun(p.add_run(text), size, bold, jp, color)
    return p


def para_runs(doc, segments, align, after=4, before=0, ls=1.3):
    """segments: [(text, size, bold, jp), ...] 混在ラン段落。"""
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = ls
    for text, size, bold, jp in segments:
        srun(p.add_run(text), size, bold, jp)
    return p


def shade(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def box(p, color="888888", sz="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom", "left", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), color)
        pbdr.append(e)
    pPr.append(pbdr)


def greeting(name):
    if is_corp(name):
        return f"{name}におかれましては、ますますご清栄のこととお慶び申し上げます。"
    return f"{name}様におかれましては、ますますご健勝のこととお慶び申し上げます。"


def write_letter(doc, rec, sender, d, headline, subject_kind):
    """建物買取DM横書き_改良版 と同じデザイン（ネイビー見出し・網掛け・囲みボックス）。"""
    bukken = f"{rec['city']}{rec['basho']}{rec['chiban']}"
    chimoku = rec["chimoku"] or "土地"
    area = rec["area"]
    given = sender["name"].split("　")[-1] if "　" in sender["name"] else sender["name"]

    # 宛先
    if rec["postal"]:
        para(doc, f"〒{rec['postal']}", 10, L, after=1)
    if rec["addr"]:
        para(doc, rec["addr"], 10.5, L, after=1)
    para(doc, f"{rec['name']}　様", 12.5, L, bold=True, after=6)

    # 見出し（ネイビー・ゴシック）
    para(doc, headline, 18, C, bold=True, after=2, ls=1.1, jp=_GOT, color=NAVY)
    para(doc, "― 仲介手数料も登記費用も、すべて弊社が負担いたします ―",
         11, C, bold=True, after=9, jp=_GOT)

    # 本文
    para(doc, "拝啓", 11, L, after=2)
    para(doc, greeting(rec["name"]) + "突然お手紙を差し上げます失礼を、まずは深くお詫び申し上げます。",
         11, JU, after=4, indent=True)
    area_txt = f"（地目：{chimoku}／地積：約{area}㎡）" if area else f"（地目：{chimoku}）"
    para(doc, f"私どもは、{subject_kind}の買取を行っております{sender['company']}の{sender['title']}、"
              f"{given}と申します。このたび、公開されている登記情報を拝見し、{bukken}の土地{area_txt}を"
              f"ご所有と拝察し、ご連絡を差し上げました。もし今後ご利用のご予定がなければ、ぜひ弊社にて"
              f"買い取らせていただきたく存じます。", 11, JU, after=4, indent=True)
    para(doc, "未活用の土地や空き家は、お使いにならなくても固定資産税や管理の負担が毎年かかり続けます。"
              "放置して「特定空家」等に指定されますと、固定資産税の軽減が外れ負担が増すこともあります。"
              "使わないうちに手放すことが、結果的に大切な資産を守る選択になるケースは少なくありません。",
         11, JU, after=8, indent=True)

    # メリット（ネイビー網掛け見出し＋●4点）
    mh = para(doc, "弊社にご売却いただく４つのメリット", 13, L, bold=True,
              after=5, before=1, jp=_GOT, color=NAVY)
    shade(mh)
    for mm in ["仲介手数料がかかりません（通常のご売却では約33万円かかります）",
               "売却登記などの諸経費（約7万円）も、すべて弊社が負担いたします",
               "広告や内覧の必要がなく、近隣に知られることなくご売却いただけます",
               "現金化までスピーディーに対応いたします"]:
        para(doc, "●　" + mm, 11.5, JU, after=3, ls=1.2)

    # 手取り強調（囲みボックス＋網掛け・数字拡大）
    kp = para_runs(doc, [
        ("通常のご売却では、仲介手数料と諸経費で ", 11.5, False, _MIN),
        ("合わせて約40万円", 15, True, _GOT),
        (" が売却代金から差し引かれます。弊社の買取なら ", 11.5, False, _MIN),
        ("これらは一切かからず", 12.5, True, _GOT),
        ("、", 11.5, False, _MIN),
        ("手取りが多く残ります。", 12.5, True, _GOT),
    ], JU, after=7, before=3, ls=1.3)
    box(kp)
    shade(kp, "F2F2F2")

    para(doc, "まずは「査定額だけ知りたい」というご相談でも構いません。お電話・FAX・メールにて、"
              "お気軽にご一報ください。無料でお見積りをお持ちいたします。どうぞよろしくお願い申し上げます。",
         11, JU, after=6, indent=True)
    para(doc, "敬具", 11, R, after=6, before=1)

    # 署名（社名はゴシック大）
    para(doc, to_reiwa(d), 10.5, R, after=2)
    addr = f"〒{sender['postal']}　{sender['address']}" if sender.get("postal") else sender["address"]
    para(doc, addr, 10.5, R, after=1)
    para(doc, sender["company"], 12.5, R, bold=True, after=1, jp=_GOT)
    para(doc, f"{sender['title']}　{sender['name']}", 10.5, R, after=1)
    tel = f"TEL {to_halfwidth(sender.get('tel',''))}"
    if sender.get("fax"):
        tel += f"　／　FAX {to_halfwidth(sender['fax'])}"
    if sender.get("mobile"):
        tel += f"　／　携帯 {to_halfwidth(sender['mobile'])}"
    para(doc, tel, 10.5, R, after=1)
    if sender.get("email"):
        para(doc, f"メール {sender['email']}", 10.5, R, after=1)


def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    sec.top_margin = sec.bottom_margin = Inches(0.5)
    sec.left_margin = sec.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = _ASCII
    normal.font.size = Pt(11)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _MIN)
    return doc


def combined_docx(recs, sender, d, headline, subject_kind):
    doc = new_doc()
    for i, rec in enumerate(recs):
        if i > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        write_letter(doc, rec, sender, d, headline, subject_kind)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


def zip_docx(recs, sender, d, headline, subject_kind):
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, rec in enumerate(recs, 1):
            doc = new_doc()
            write_letter(doc, rec, sender, d, headline, subject_kind)
            b = io.BytesIO()
            doc.save(b)
            safe = "".join(ch for ch in rec["name"] if ch not in '\\/:*?"<>|').strip()
            zf.writestr(f"{i:03d}_{safe}.docx", b.getvalue())
    buf.seek(0)
    return buf


# ─────────────────────────── UI ───────────────────────────
st.set_page_config(page_title="買取DMメーカー", page_icon="🏠", layout="wide")
st.title("🏠 買取DMメーカー")
st.caption("所有者台帳 → 未活用地・空き家の買取DMを差し込み量産（謄本から台帳へ追加も可）")

if "daicho" not in st.session_state:
    st.session_state["daicho"] = None

senders = load_senders()

# ─── サイドバー：差出人（追加・編集） ───
with st.sidebar:
    st.header("差出人")
    labels = [s["label"] for s in senders]
    sel = st.selectbox("この差出人でDMを作成", labels, index=0)
    sender = next(s for s in senders if s["label"] == sel)
    st.caption(f"{sender['company']}　{sender['title']} {sender['name']}")

    with st.expander("差出人を追加・編集", expanded=False):
        edit_target = st.selectbox("編集対象", ["＋ 新規追加"] + labels, index=0, key="edit_target")
        base = {} if edit_target == "＋ 新規追加" else next(s for s in senders if s["label"] == edit_target)
        vals = {}
        for key, jp in SENDER_FIELDS:
            vals[key] = st.text_input(jp, base.get(key, ""), key=f"sf_{key}")
        cc1, cc2 = st.columns(2)
        with cc1:
            if st.button("💾 保存", use_container_width=True):
                if not vals["label"].strip():
                    st.warning("表示名（ラベル）は必須です。")
                else:
                    if edit_target == "＋ 新規追加":
                        senders.append(vals)
                    else:
                        i = next(i for i, s in enumerate(senders) if s["label"] == edit_target)
                        senders[i] = vals
                    save_senders(senders)
                    st.success("保存しました。")
                    st.rerun()
        with cc2:
            if edit_target != "＋ 新規追加" and st.button("🗑 削除", use_container_width=True):
                senders = [s for s in senders if s["label"] != edit_target]
                save_senders(senders)
                st.success("削除しました。")
                st.rerun()

    st.divider()
    doc_date = st.date_input("送付日", value=date.today())

    # ─── サイドバー：台帳更新 ───
    st.divider()
    st.header("台帳更新")
    up = st.file_uploader("① 台帳をアップロード（.xls/.xlsx）", type=["xls", "xlsx"])
    if up is not None and st.button("台帳を読み込む", use_container_width=True):
        try:
            xls = pd.ExcelFile(up)
            sh = "NO順" if "NO順" in xls.sheet_names else xls.sheet_names[0]
            df = xls.parse(sh)
            df.columns = [str(c).strip() for c in df.columns]
            st.session_state["daicho"] = df
            st.success(f"台帳を読み込みました（{len(df)}行）。")
            st.rerun()
        except Exception as e:
            st.error(f"読み込みエラー: {e}")

    st.markdown("**② 謄本を台帳に追加**")
    touhon = st.file_uploader("謄本PDF（複数可・5件程度まで）",
                              type=["pdf"], accept_multiple_files=True, key="touhon")
    merge_mode = st.radio(
        "取り込み方",
        ["1ファイル＝1物件（まとめて追加）", "全ファイル＝1物件に統合（土地＋建物が別ファイル）"],
        index=0, key="merge_mode")
    if touhon:
        st.caption(f"アップロード：{len(touhon)} ファイル")
    if touhon and st.button("🧾 謄本を読み取って台帳に追加", type="primary", use_container_width=True):
        try:
            rp = _get_registry_parser()
        except Exception as e:
            st.error(f"謄本パーサを読み込めません: {e}")
            rp = None
        if rp is not None:
            def wrap(f):
                b = io.BytesIO(f.read())
                b.name = f.name
                return b
            cur = st.session_state["daicho"]
            if cur is None or cur.empty:
                cur = pd.DataFrame(columns=CANON_COLS)
            new_rows = []
            if merge_mode.startswith("全ファイル"):
                with st.spinner(f"謄本をAIで読み取り中…（{len(touhon)}ファイルを1物件に統合）"):
                    parsed = rp.parse_registry([wrap(f) for f in touhon])
                new_rows.append(registry_to_record(parsed, next_no=len(cur) + 1))
            else:
                prog = st.progress(0.0)
                for i, f in enumerate(touhon):
                    with st.spinner(f"読み取り中… {i + 1}/{len(touhon)}：{f.name}"):
                        parsed = rp.parse_registry([wrap(f)])
                    new_rows.append(registry_to_record(parsed, next_no=len(cur) + len(new_rows) + 1))
                    prog.progress((i + 1) / len(touhon))
            st.session_state["daicho"] = pd.concat(
                [cur, pd.DataFrame(new_rows)], ignore_index=True)
            st.success(f"台帳に {len(new_rows)} 行追加しました。")
            st.rerun()

    if st.session_state["daicho"] is not None and not st.session_state["daicho"].empty:
        st.download_button(
            "⬇️ 更新した台帳をダウンロード",
            daicho_to_xlsx(st.session_state["daicho"]),
            file_name="台帳.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True)

# ─── メイン：台帳の確認とDM生成 ───
df = st.session_state["daicho"]
if df is None or df.empty:
    st.info("左の「台帳更新」から台帳をアップロード、または謄本を追加してください。")
    st.stop()

st.subheader("台帳")
st.dataframe(df, use_container_width=True, height=260)

m = auto_map(df.columns)
if not m.get("name"):
    st.error("『登記名義人』列が見つかりません。台帳の列名をご確認ください。")
    st.stop()

st.subheader("絞り込み")
f1, f2, f3, f4 = st.columns(4)
with f1:
    require_addr = st.checkbox("現住所なしを除外", value=True)
with f2:
    dedup = st.checkbox("宛先重複を集約", value=True)
with f3:
    who = st.radio("対象", ["すべて", "個人のみ", "法人のみ"], index=0)
with f4:
    excl_mortgage = st.checkbox("備考に抵当権ありを除外", value=False)

recs = build_records(df, m, require_addr=require_addr)
if who == "個人のみ":
    recs = [r for r in recs if not is_corp(r["name"])]
elif who == "法人のみ":
    recs = [r for r in recs if is_corp(r["name"])]
if excl_mortgage:
    recs = [r for r in recs if "抵当" not in r["biko"]]
recs = dedupe(recs) if dedup else [dict(r, extra=0) for r in recs]

st.success(f"送付対象：**{len(recs)} 通**（台帳 {len(df)} 行）")

if recs:
    master = st.checkbox("すべて選択（外すと全解除から選べます）", value=True, key="master_sel")
    prev = pd.DataFrame([{
        "送付": master,
        "名義人": r["name"],
        "種別": "法人" if is_corp(r["name"]) else "個人",
        "物件": f"{r['city']}{r['basho']}{r['chiban']}",
        "地目": r["chimoku"], "地積㎡": r["area"],
        "送付先": f"〒{r['postal']} {r['addr']}".strip(),
        "他筆": r["extra"] or "",
    } for r in recs])
    edited = st.data_editor(
        prev, use_container_width=True, height=320, hide_index=True,
        key=f"editor_{master}",
        column_config={"送付": st.column_config.CheckboxColumn("送付", default=True)},
        disabled=[c for c in prev.columns if c != "送付"],
    )
    keep = list(edited["送付"])
    target = [r for r, k in zip(recs, keep) if k]

    st.caption(f"選択：**{len(target)} / {len(recs)} 通**")
    mode = st.radio("出力形式", ["結合docx（1通1ページ）", "名義人ごと個別docx（ZIP）"], index=0)

    if st.button("📄 チェックした宛先にDMを生成", type="primary", disabled=not target):
        if mode.startswith("結合"):
            data = combined_docx(target, sender, doc_date, HEADLINE, SUBJECT_KIND)
            st.download_button("⬇️ 結合docxをダウンロード", data,
                               file_name=f"買取DM_{len(target)}通.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            data = zip_docx(target, sender, doc_date, HEADLINE, SUBJECT_KIND)
            st.download_button("⬇️ ZIP（個別docx）をダウンロード", data,
                               file_name=f"買取DM_{len(target)}通.zip",
                               mime="application/zip")
        st.success(f"{len(target)} 通を生成しました。")
