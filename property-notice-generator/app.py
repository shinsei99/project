import streamlit as st
import subprocess
import json
import os
import html
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CLAUDE_BIN = os.path.expanduser("~/.local/bin/claude")
CLAUDE_TIMEOUT = 120

# ── Constants ────────────────────────────────────────────────────────────────
COMPANY_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "company_info.json")

PURPOSE_OPTIONS = [
    # 生活マナー系
    "騒音注意",
    "ゴミ出しマナー",
    "喫煙マナー（共用部・ベランダ等）",
    "ペットの飼育・マナー",
    "洗濯物・布団干しのマナー",
    "エレベーターの利用マナー",
    # 共用スペース系
    "違法駐車・駐輪",
    "共有部の私物放置",
    "自転車・バイクの整理整頓",
    "宅配ボックスの利用マナー",
    "駐車場の利用ルール",
    # 設備・工事系
    "設備点検のお知らせ",
    "工事・作業のお知らせ（断水・停電等）",
    "共用設備の使用停止・変更",
    # 安全・防犯系
    "防火・防犯へのご協力",
    "不審者・セキュリティに関するお知らせ",
    "台風・災害時の注意事項",
    # 契約・管理系
    "家賃・管理費の支払いについて",
    "契約更新のご案内",
    "退去・明け渡し手続きのご案内",
    "管理規約・ルール改定のお知らせ",
    # 自由入力
    "その他（自由入力）",
]
TARGET_OPTIONS = [
    "入居者の皆様へ",
    "特定の階・お部屋の方へ",
    "特定の箇所の利用者様へ",
]
TONE_OPTIONS = ["マイルドに注意", "厳しく警告", "丁寧にお願い", "重要なお知らせ"]

TITLE_MAP = {
    # 生活マナー系
    "騒音注意": "騒音に関するお願い",
    "ゴミ出しマナー": "ゴミの分別・出し方についてのお願い",
    "喫煙マナー（共用部・ベランダ等）": "喫煙マナーに関するお願い",
    "ペットの飼育・マナー": "ペットの飼育・マナーに関するお願い",
    "洗濯物・布団干しのマナー": "洗濯物・布団干しに関するお願い",
    "エレベーターの利用マナー": "エレベーター利用マナーに関するお願い",
    # 共用スペース系
    "違法駐車・駐輪": "駐車・駐輪についてのお願い",
    "共有部の私物放置": "共用部分の使用についてのお願い",
    "自転車・バイクの整理整頓": "自転車・バイクの整理整頓についてのお願い",
    "宅配ボックスの利用マナー": "宅配ボックスの利用マナーについて",
    "駐車場の利用ルール": "駐車場の利用ルールについて",
    # 設備・工事系
    "設備点検のお知らせ": "設備点検のお知らせ",
    "工事・作業のお知らせ（断水・停電等）": "工事・作業に関するお知らせ",
    "共用設備の使用停止・変更": "共用設備の使用変更に関するお知らせ",
    # 安全・防犯系
    "防火・防犯へのご協力": "防火・防犯へのご協力のお願い",
    "不審者・セキュリティに関するお知らせ": "防犯・セキュリティに関するお知らせ",
    "台風・災害時の注意事項": "台風・災害時の注意事項",
    # 契約・管理系
    "家賃・管理費の支払いについて": "家賃・管理費のお支払いについて",
    "契約更新のご案内": "契約更新のご案内",
    "退去・明け渡し手続きのご案内": "退去・明け渡し手続きのご案内",
    "管理規約・ルール改定のお知らせ": "管理規約改定のお知らせ",
    # 自由入力
    "その他（自由入力）": "お知らせ・お願い",
}
TONE_PREFIX = {
    "マイルドに注意": "",
    "厳しく警告": "【警告】",
    "丁寧にお願い": "",
    "重要なお知らせ": "【重要】",
}

# ── Persistence ──────────────────────────────────────────────────────────────
def load_company_info():
    if os.path.exists(COMPANY_FILE):
        try:
            with open(COMPANY_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"name": "", "address": "", "contact_person": "", "phone": ""}


def save_company_info(name: str, address: str, contact_person: str, phone: str) -> dict:
    data = {"name": name, "address": address, "contact_person": contact_person, "phone": phone}
    with open(COMPANY_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return data


# ── AI Generation (Claude Code CLI) ─────────────────────────────────────────
def generate_notice(
    purpose: str,
    custom_purpose: str,
    target: str,
    situation: str,
    tone: str,
    company_name: str,
) -> str:
    actual_purpose = custom_purpose if purpose == "その他（自由入力）" else purpose

    tone_guide = {
        "マイルドに注意": "穏やかで柔らかい表現を使い、相手を傷つけずに改善を促す口調。",
        "厳しく警告": "毅然とした警告口調。問題の深刻さを明示し、改善がない場合の管理規約に基づく措置についても言及する。",
        "丁寧にお願い": "最大限丁寧な敬語を使用し、感謝の意も示しながら協力を心からお願いする口調。",
        "重要なお知らせ": "客観的かつ明確に情報を伝え、必要なアクションや注意点を箇条書き等で整理して示す。",
    }

    prompt = f"""不動産管理会社の熟練した文書担当者として、入居者向けの通知文（本文のみ）を作成してください。

■ 文書の目的: {actual_purpose}
■ 対象: {target}
■ 具体的な状況・詳細: {situation.strip() if situation.strip() else "（詳細未入力）"}
■ トーン: {tone}（{tone_guide.get(tone, "")}）
■ 管理会社名: {company_name or "管理会社"}

【作成ルール】
- 本文のみを出力すること（タイトル・日付・署名・連絡先は絶対に含めない）
- 書き出しは「平素より当物件をご利用いただき、誠にありがとうございます。」など自然な形で始める
- 2〜4段落で適切に段落分けする（段落間は空行を入れる）
- 必要に応じて箇条書き（「・」を使用）を活用する
- 締めくくりは「引き続きご理解とご協力のほど、よろしくお願い申し上げます。」で終える
- 分量はA4一枚に収まる400〜600文字程度
- 管理会社として誠実で専門的な印象を与える文体にすること
- 角が立たず、かつ伝えるべきことはしっかり伝えること

本文のみを出力してください:"""

    cmd = [
        CLAUDE_BIN, "-p", prompt,
        "--output-format", "json",
        "--dangerously-skip-permissions",
        "--model", "sonnet",
    ]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=CLAUDE_TIMEOUT)
    except FileNotFoundError:
        raise RuntimeError("`claude` コマンドが見つかりません。Claude Code CLI がインストールされているか確認してください。")
    except subprocess.TimeoutExpired:
        raise RuntimeError(f"生成が{CLAUDE_TIMEOUT}秒を超えたため中断しました。再試行してください。")

    if proc.returncode != 0:
        raise RuntimeError(f"claude コマンドが失敗しました（終了コード {proc.returncode}）\n{proc.stderr.strip()[:300]}")

    result = json.loads(proc.stdout)
    if result.get("is_error"):
        raise RuntimeError(f"Claude がエラーを返しました: {result.get('result')}")
    return result.get("result", "").strip()


# ── Word Document Generation ─────────────────────────────────────────────────
def _set_font(run, size_pt: float, bold: bool = False, color_rgb=None, font: str = "游ゴシック"):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)


def _add_hline(doc, color: str = "1B4B8A", sz: str = "12"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single")
    btm.set(qn("w:sz"), sz)
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), color)
    pBdr.append(btm)
    pPr.append(pBdr)


def _spacer(doc, after_pt: float = 4.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after_pt)


def create_docx(doc_title: str, body_text: str, target: str, cinfo: dict) -> BytesIO:
    doc = Document()

    # A4 page setup
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)

    # ── Title with blue background ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(14)
    pPr = title_p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1B4B8A")
    pPr.append(shd)
    tr = title_p.add_run(f"\n{doc_title}\n")
    _set_font(tr, 18, bold=True, color_rgb=(255, 255, 255))

    # ── Date (right) ──
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dp.paragraph_format.space_after = Pt(2)
    _set_font(dp.add_run("令和　　年　　月　　日"), 10.5)

    # ── Target (left, bold) ──
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(2)
    _set_font(tp.add_run(target), 11, bold=True)

    # ── Company info (right) ──
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cp.paragraph_format.space_after = Pt(8)
    lines = [cinfo.get("name", "")]
    if cinfo.get("address"):
        lines.append(cinfo["address"])
    if cinfo.get("contact_person"):
        lines.append(f"担当：{cinfo['contact_person']}")
    if cinfo.get("phone"):
        lines.append(f"TEL：{cinfo['phone']}")
    _set_font(cp.add_run("\n".join(lines)), 10)

    # ── Horizontal rule ──
    _add_hline(doc)
    _spacer(doc, 4)

    # ── Body text ──
    for line in body_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        if line.startswith(("・", "•", "●", "▶", "◆")):
            p.paragraph_format.left_indent = Cm(0.5)
        else:
            p.paragraph_format.first_line_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(5)
        _set_font(p.add_run(line), 10.5)

    _spacer(doc, 6)

    # ── Bottom horizontal rule ──
    _add_hline(doc)

    # ── Footer contact ──
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(5)
    parts = ["【お問い合わせ先】　", cinfo.get("name", "")]
    if cinfo.get("address"):
        parts.append(f"　〒{cinfo['address']}" if not cinfo["address"].startswith("〒") else f"　{cinfo['address']}")
    if cinfo.get("contact_person"):
        parts.append(f"　担当：{cinfo['contact_person']}")
    if cinfo.get("phone"):
        parts.append(f"　TEL：{cinfo['phone']}")
    _set_font(fp.add_run("".join(parts)), 10, bold=True, color_rgb=(27, 75, 138))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Preview HTML ─────────────────────────────────────────────────────────────
def render_preview(doc_title: str, body_text: str, target: str, cinfo: dict):
    e = html.escape

    company_html = e(cinfo.get("name", ""))
    if cinfo.get("address"):
        company_html += "<br>" + e(cinfo["address"])
    if cinfo.get("contact_person"):
        company_html += "<br>" + e(f"担当：{cinfo['contact_person']}")
    if cinfo.get("phone"):
        company_html += "<br>" + e(f"TEL：{cinfo['phone']}")

    footer_parts = [e(cinfo.get("name", ""))]
    if cinfo.get("address"):
        addr = cinfo["address"]
        footer_parts.append(e(f"〒{addr}" if not addr.startswith("〒") else addr))
    if cinfo.get("contact_person"):
        footer_parts.append(e(f"担当：{cinfo['contact_person']}"))
    if cinfo.get("phone"):
        footer_parts.append(e(f"TEL：{cinfo['phone']}"))
    footer_str = "　".join(footer_parts)

    body_html = e(body_text).replace("\n", "<br>")

    st.markdown(
        f"""
<div style="
    background:white;border:1px solid #ccc;border-radius:8px;padding:0;
    font-family:'Hiragino Sans','Yu Gothic','MS Gothic',sans-serif;
    color:#222;box-shadow:0 3px 12px rgba(0,0,0,0.12);max-width:680px;margin:0 auto;">
  <div style="
      background:#1B4B8A;color:white;text-align:center;
      padding:14px 24px;font-size:17px;font-weight:bold;
      border-radius:8px 8px 0 0;line-height:1.5;">
    {e(doc_title)}
  </div>
  <div style="padding:20px 30px 24px 30px;">
    <div style="text-align:right;font-size:12px;color:#666;margin-bottom:8px;">
      令和　　年　　月　　日
    </div>
    <div style="display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:10px;">
      <div style="font-weight:bold;font-size:14px;">{e(target)}</div>
      <div style="text-align:right;font-size:12px;color:#444;line-height:1.8;">{company_html}</div>
    </div>
    <hr style="border:none;border-top:2.5px solid #1B4B8A;margin:8px 0 16px 0;">
    <div style="font-size:14px;line-height:2.1;text-indent:1em;">{body_html}</div>
    <hr style="border:none;border-top:2.5px solid #1B4B8A;margin:16px 0 10px 0;">
    <div style="text-align:center;font-size:13px;font-weight:bold;color:#1B4B8A;">
      【お問い合わせ先】　{footer_str}
    </div>
  </div>
</div>
""",
        unsafe_allow_html=True,
    )


# ── Main App ─────────────────────────────────────────────────────────────────
def main():
    st.set_page_config(
        page_title="物件管理 案内文ジェネレーター",
        page_icon="🏢",
        layout="wide",
    )

    if "company_info" not in st.session_state:
        st.session_state.company_info = load_company_info()

    # ── Sidebar ──────────────────────────────────────────────────────────────
    with st.sidebar:
        st.markdown("## 🏢 管理会社情報")
        st.caption("入力後「保存」を押すと次回も自動入力されます")

        ci = st.session_state.company_info
        c_name = st.text_input("会社名", value=ci.get("name", ""), placeholder="例：〇〇不動産管理株式会社")
        c_address = st.text_input("住所", value=ci.get("address", ""), placeholder="例：東京都新宿区〇〇1-2-3")
        c_person = st.text_input("担当者名", value=ci.get("contact_person", ""), placeholder="例：田中 太郎")
        c_phone = st.text_input("電話番号", value=ci.get("phone", ""), placeholder="例：03-1234-5678")

        if st.button("💾 管理会社情報を保存", use_container_width=True, type="secondary"):
            st.session_state.company_info = save_company_info(c_name, c_address, c_person, c_phone)
            st.success("✅ 保存しました！")

    current_cinfo = {"name": c_name, "address": c_address, "contact_person": c_person, "phone": c_phone}

    # ── Main area ────────────────────────────────────────────────────────────
    st.title("🏢 物件管理 案内文ジェネレーター")
    st.caption("状況を選ぶだけでAIが案内文を自動生成 → 印刷用Wordファイル（.docx）をダウンロード")
    st.divider()

    left, right = st.columns([1, 1.15], gap="large")

    with left:
        st.subheader("📝 文書の内容を入力")

        purpose = st.selectbox(
            "① 文書の目的",
            PURPOSE_OPTIONS,
            format_func=lambda x: (
                f"── {x[2:]} ──" if x.startswith("# ") else x
            ),
        )

        custom_purpose = ""
        if purpose == "その他（自由入力）":
            custom_purpose = st.text_input(
                "目的を具体的に入力してください",
                placeholder="例：共用廊下での喫煙について",
            )

        target_type = st.selectbox("② 対象", TARGET_OPTIONS)

        target_detail = ""
        if target_type != "入居者の皆様へ":
            target_detail = st.text_input(
                "対象の詳細を入力してください",
                placeholder="例：3階 301号室の方へ　/　自転車置き場をご利用の方へ",
            )

        situation = st.text_area(
            "③ 具体的な状況",
            placeholder=(
                "例：夜間23時以降のドアの開閉音や足音が大きく、\n"
                "他の入居者様よりご苦情をいただいております。"
            ),
            height=130,
        )

        tone = st.selectbox("④ 文書のトーン", TONE_OPTIONS)

        tone_desc = {
            "マイルドに注意": "🌿 穏やかに改善をお願いするスタイル",
            "厳しく警告": "⚠️ 毅然とした警告。規約違反の措置にも言及",
            "丁寧にお願い": "🙏 最大限丁寧な敬語で協力を依頼",
            "重要なお知らせ": "📢 客観的・明確に情報を伝えるスタイル",
        }
        st.caption(tone_desc.get(tone, ""))

        final_target = target_detail.strip() if target_detail.strip() else target_type

        gen_btn = st.button("✨ 案内文を生成する", type="primary", use_container_width=True)

    # ── Right column ─────────────────────────────────────────────────────────
    with right:
        st.subheader("📄 プレビュー・ダウンロード")

        if gen_btn:
            if purpose == "その他（自由入力）" and not custom_purpose.strip():
                st.warning("⚠️ 文書の目的を入力してください。")
            else:
                with st.spinner("AIが案内文を生成しています..."):
                    try:
                        body = generate_notice(
                            purpose,
                            custom_purpose,
                            final_target,
                            situation,
                            tone,
                            c_name,
                        )
                        st.session_state["body_text_area"] = body
                        st.session_state["has_content"] = True
                        st.session_state["doc_meta"] = {
                            "doc_title": TONE_PREFIX[tone] + TITLE_MAP.get(purpose, "お知らせ"),
                            "final_target": final_target,
                            "cinfo": current_cinfo,
                        }
                        st.success("✅ 生成完了！本文を確認・編集してダウンロードしてください。")
                    except Exception as ex:
                        st.error(f"❌ エラーが発生しました: {ex}")

        if st.session_state.get("has_content"):
            meta = st.session_state["doc_meta"]
            doc_title = meta["doc_title"]
            ft = meta["final_target"]
            # Use current sidebar values for cinfo (user may update company info)
            cinfo_for_doc = current_cinfo

            # ── Editable text area ──
            edited = st.text_area(
                "✏️ 本文（編集可能）",
                height=220,
                key="body_text_area",
                help="AIが生成した本文を自由に編集できます。",
            )

            st.markdown("---")

            # ── Document preview ──
            with st.expander("📋 完成イメージ（プレビュー）", expanded=True):
                render_preview(doc_title, edited, ft, cinfo_for_doc)

            st.markdown("---")

            # ── Download button ──
            try:
                word_buf = create_docx(doc_title, edited, ft, cinfo_for_doc)
                fname = f"notice_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.download_button(
                    label="📥 配布用Word（.docx）をダウンロード",
                    data=word_buf,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary",
                )
                st.caption(f"A4印刷用フォーマット・ファイル名: {fname}")
            except Exception as ex:
                st.error(f"Word生成エラー: {ex}")

        elif not gen_btn:
            st.info("👈 左側で内容を入力し、「案内文を生成する」ボタンを押してください。")


if __name__ == "__main__":
    main()
