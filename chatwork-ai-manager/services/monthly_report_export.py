"""業務月報の書き出し（Word / Excel）と Chatwork 投稿用の整形。

月報は日報と違って「1人1本」ではなく「1回のトリガーにつき1本」なので、
日報のような1シート内の複数人ループは無い。見出し記法（`## X` + `- 箇条書き`）は
日報と共通なので、解釈（parse_sections）と罫線・フォント設定は daily_report_export を再利用する。
"""
import datetime

from services.daily_report_export import _set_font, _table_borders, _wrapped_lines, parse_sections

_WD = "月火水木金土日"


def period_label(period: str) -> str:
    try:
        y, m = period.split("-")
        return f"{y}年{int(m)}月"
    except ValueError:
        return period


# ---- Chatwork 投稿用 ---------------------------------------------------------
def chatwork_body(row) -> str:
    out = [f"📝 業務月報（{period_label(row['report_period'])}・AI作成のたたき台）", ""]
    for heading, lines in parse_sections(row["body"]):
        out.append(f"【{heading}】")
        out += [f"・{ln}" for ln in lines] or ["・特になし"]
        out.append("")
    return "\n".join(out).strip()


# ---- Word --------------------------------------------------------------------
def build_docx(row, path: str) -> str:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(2.0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run("業 務 月 報"), size=18, bold=True)

    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _table_borders(t)
    _set_font(t.cell(0, 0).paragraphs[0].add_run("対象月"), size=10, bold=True)
    _set_font(t.cell(0, 1).paragraphs[0].add_run(period_label(row["report_period"])), size=10)
    _set_font(t.cell(1, 0).paragraphs[0].add_run("要　約"), size=10, bold=True)
    _set_font(t.cell(1, 1).paragraphs[0].add_run(row["summary"] or "-"), size=10)
    for w, col in ((Cm(2.2), 0), (Cm(10.2), 1)):
        for cell in t.columns[col].cells:
            cell.width = w

    doc.add_paragraph()
    for heading, lines in parse_sections(row["body"]):
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(10)
        hp.paragraph_format.space_after = Pt(2)
        _set_font(hp.add_run(f"■ {heading}"), size=11, bold=True)
        for ln in (lines or ["特になし"]):
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(0)
            bp.paragraph_format.left_indent = Cm(1.1)
            bp.paragraph_format.first_line_indent = Cm(-0.55)
            _set_font(bp.add_run(ln), size=10)

    fp = doc.add_paragraph()
    fp.paragraph_format.space_before = Pt(14)
    r = fp.add_run(f"生成: {row['model'] or '-'}（{row['generated_by'] or '-'}）")
    _set_font(r, size=8)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(path)
    return path


# ---- Excel -------------------------------------------------------------------
def sheet_name(period: str) -> str:
    return f"業務月報{period.replace('-', '')}"


def build_xlsx(row, path: str) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    thin = Side(style="thin", color="AAAAAA")
    box = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill("solid", fgColor="E8EEF7")
    title_font = Font(name="游ゴシック", size=16, bold=True)
    label_font = Font(name="游ゴシック", size=10, bold=True)
    body_font = Font(name="游ゴシック", size=10)

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name(row["report_period"])
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 74

    ws["A1"] = "業務月報"
    ws["A1"].font = title_font
    ws.merge_cells("A1:B1")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    for r, (label, value) in enumerate(
            (("対象月", period_label(row["report_period"])), ("要約", row["summary"] or "-")), start=2):
        ws.cell(r, 1, label).font = label_font
        ws.cell(r, 1).fill = head_fill
        ws.cell(r, 2, value).font = body_font
        for c in (1, 2):
            ws.cell(r, c).border = box
            ws.cell(r, c).alignment = Alignment(vertical="center", wrap_text=True)

    r = 5
    for heading, lines in parse_sections(row["body"]):
        ws.cell(r, 1, heading).font = label_font
        ws.cell(r, 1).fill = head_fill
        text = "\n".join(("特になし" if ln == "特になし" else f"・{ln}")
                         for ln in (lines or ["特になし"]))
        ws.cell(r, 2, text).font = body_font
        for c in (1, 2):
            ws.cell(r, c).border = box
            ws.cell(r, c).alignment = Alignment(vertical="top", wrap_text=True)
        n = _wrapped_lines(text, 74)
        ws.row_dimensions[r].height = max(18, min(600, n * 18))
        r += 1

    ws.sheet_view.showGridLines = False
    ws.print_options.horizontalCentered = True
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToWidth = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    wb.save(path)
    return path
