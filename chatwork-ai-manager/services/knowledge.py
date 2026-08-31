"""社内ナレッジの取込: ファイル → テキスト抽出 → チャンク → FTS5 索引。

対応形式: PDF / Excel(xlsx) / Word(docx) / CSV / txt・md。
それ以外（画像等）はスキップ（将来 claude vision で OCR 予定）。
source_ref に「ファイル名 / P<ページ> / Sheet:<シート>」を残し、Q&Aで根拠提示できるようにする。
バージョン管理: 同じ (category, title) を再取込すると version+1・旧版は active=0（§23）。
"""
import csv
import datetime
import hashlib
import os
import unicodedata

from db.connection import get_conn, query, query_one

CHUNK_SIZE = 800          # 1 チャンクの目安文字数
CHUNK_OVERLAP = 100
SUPPORTED_EXT = {".pdf", ".xlsx", ".xlsm", ".docx", ".csv", ".txt", ".md", ".html", ".htm", ".url",
                 # ★旧Office形式（2026-08-31 追加）。取引・契約フォルダの2,418件のうち
                 #   .xls が353件・.doc が251件あり、契約書・請求書・賃貸資料など中身が
                 #   テキストの重要書類が読めないままだった。
                 ".xls", ".doc"}
SKIP_DIRS = set()                                 # 除外フォルダ（現状なし）
SKIP_DIR_PREFIXES = ("_アーカイブ", "_bak", ".")   # アーカイブ・バックアップ・隠しは除外


def _skip_dir(name: str) -> bool:
    # macOS のファイル名は NFD で返るため NFC 正規化してから比較する
    n = unicodedata.normalize("NFC", name)
    if n in SKIP_DIRS:
        return True
    return any(n.startswith(p) for p in SKIP_DIR_PREFIXES)


# ---- テキスト抽出（形式別。戻り値: [(text, source_ref), ...]） ----
def _extract_pdf(path):
    import fitz  # pymupdf
    out = []
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, 1):
            text = page.get_text("text").strip()
            if text:
                out.append((text, f"P{i}"))
    return out


def _extract_xlsx(path):
    import openpyxl
    out = []
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for ws in wb.worksheets:
        lines = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c not in (None, "")]
            if cells:
                lines.append(" | ".join(cells))
        if lines:
            out.append(("\n".join(lines), f"Sheet:{ws.title}"))
    wb.close()
    return out


def _extract_docx(path):
    import docx
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    text = "\n".join(paras).strip()
    return [(text, None)] if text else []


def _extract_xls(path):
    """旧Excel（.xls）。xlrd で読む。

    ★xlrd 2.x は .xlsx を捨てて **.xls 専用**になった。新しい方は openpyxl が見るので、
      役割がきれいに分かれている。
    """
    import xlrd
    out = []
    book = xlrd.open_workbook(path, on_demand=True)
    for sh in book.sheets():
        lines = []
        for r in range(sh.nrows):
            cells = []
            for c in sh.row(r):
                v = c.value
                if v is None or v == "":
                    continue
                # 数値が 1234.0 のように出るのを整える
                if isinstance(v, float) and v == int(v):
                    v = int(v)
                cells.append(str(v).strip())
            if cells:
                lines.append(" | ".join(cells))
        if lines:
            out.append(("\n".join(lines), f"Sheet:{sh.name}"))
    book.release_resources()
    return out


def _extract_doc(path):
    """旧Word（.doc）。macOS 標準の textutil で取り出す。

    ★antiword や olefile は入っていない。textutil は OS 同梱で、実データで試すと
      送付書・解約通知書などは問題なく読めた（1,000字前後）。
      中身が図だけの書類（領収書のひな形など）は0文字になるが、それは正しい結果。
    """
    import subprocess
    try:
        r = subprocess.run(["textutil", "-convert", "txt", "-stdout", path],
                           capture_output=True, timeout=60)
    except Exception:
        return []
    text = r.stdout.decode("utf-8", "ignore").strip()
    if not text:
        # textutil は Shift_JIS のまま返すことがある
        text = r.stdout.decode("cp932", "ignore").strip()
    return [(text, None)] if text else []


def _extract_csv(path):
    lines = []
    for enc in ("utf-8-sig", "cp932", "utf-8"):
        try:
            with open(path, newline="", encoding=enc) as f:
                for row in csv.reader(f):
                    cells = [c for c in row if c.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            break
        except (UnicodeDecodeError, UnicodeError):
            lines = []
            continue
    return [("\n".join(lines), None)] if lines else []


def _extract_text(path):
    for enc in ("utf-8", "cp932"):
        try:
            with open(path, encoding=enc) as f:
                t = f.read().strip()
            return [(t, None)] if t else []
        except (UnicodeDecodeError, UnicodeError):
            continue
    return []


class _HTMLSections(__import__("html.parser", fromlist=["HTMLParser"]).HTMLParser):
    """HTML を <section> 単位に分け、各セクションの見出しを source_ref にする。"""
    _SKIP = {"script", "style", "noscript"}
    _HEAD = {"h1", "h2", "h3", "h4"}

    def __init__(self):
        super().__init__()
        self.sections = []
        self.cur, self.cur_title = [], None
        self._skip, self._in_head = 0, False

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip += 1
        elif tag == "section":
            self._flush()
        elif tag in self._HEAD:
            self._in_head = True

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip:
            self._skip -= 1
        elif tag in self._HEAD:
            self._in_head = False

    def handle_data(self, data):
        if self._skip:
            return
        t = data.strip()
        if not t:
            return
        if self._in_head and self.cur_title is None:
            self.cur_title = t
        self.cur.append(t)

    def _flush(self):
        text = "\n".join(self.cur).strip()
        if text:
            self.sections.append((text, self.cur_title))
        self.cur, self.cur_title = [], None

    def result(self):
        self._flush()
        return self.sections


def _extract_html(path):
    for enc in ("utf-8", "cp932"):
        try:
            with open(path, encoding=enc) as f:
                html = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            html = None
    if not html:
        return []
    p = _HTMLSections()
    p.feed(html)
    return p.result()


def _extract_url(path):
    """Windows .url ショートカットから URL を取り出し、社内ツール案内用テキストにする。"""
    url = None
    for enc in ("utf-8", "cp932", "latin-1"):
        try:
            with open(path, encoding=enc, errors="ignore") as f:
                for line in f:
                    if line.strip().lower().startswith("url="):
                        url = line.strip()[4:].strip()
                        break
            break
        except Exception:
            continue
    if not url:
        return []
    name = os.path.splitext(os.path.basename(path))[0]
    text = (f"社内ツール（社内Webアプリ）: {name}\n"
            f"アクセスURL: {url}\n"
            f"社内LANのブラウザで上記URLを開いて使う社内システム。")
    return [(text, "社内ツール")]


_EXTRACTORS = {
    ".pdf": _extract_pdf, ".xlsx": _extract_xlsx, ".xlsm": _extract_xlsx,
    ".docx": _extract_docx, ".csv": _extract_csv, ".txt": _extract_text, ".md": _extract_text,
    ".html": _extract_html, ".htm": _extract_html, ".url": _extract_url,
    # 旧Office（2026-08-31 追加）。SUPPORTED_EXT に足すだけでは読めない。
    # extract() はこの表を引くので、**両方に足す**こと。
    ".xls": _extract_xls, ".doc": _extract_doc,
}


def extract(path):
    ext = os.path.splitext(path)[1].lower()
    fn = _EXTRACTORS.get(ext)
    if not fn:
        return []
    try:
        return fn(path)
    except Exception as e:
        raise RuntimeError(f"抽出失敗 {os.path.basename(path)}: {type(e).__name__}: {e}")


def _render_pdf_images(path, out_dir, max_pages=15, dpi=150):
    import fitz
    out = []
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, 1):
            if i > max_pages:
                break
            pix = page.get_pixmap(dpi=dpi)
            fp = os.path.join(out_dir, f"page-{i:02d}.png")
            pix.save(fp)
            out.append((i, fp))
    return out


class OcrSkippedByQuotaSaver(Exception):
    """節約モード中なので claude へ回さなかった。**異常ではない**（2026-09-01）。

    ★`OcrUnavailable` と分けるのが肝。`OcrUnavailable` は「claudeが壊れている」の合図で、
      呼び出し側は連続3回でその晩を諦める作りになっている。節約モードでこれを投げたら、
      **Visionで読めない書類に3件当たっただけで一晩が終わった**（2026-09-01の実測: 15分・5件で撤退）。
      こちらは「この書類は後日に回す」だけの意味で、撤退の理由にはしない。
    """


class OcrUnavailable(Exception):
    """OCRの土台（claude）が使えなかった。**書類に文字が無いのとは別物**。

    これを「文字が無い」と同じ扱いにすると、定額枠が切れていた日に当たった書類が
    見送りリストへ入り、**二度とOCRされなくなる**（2026-08-28に実際に発生）。
    """


# ★macOS Vision の文字起こし（2026-08-31 に二段構えへ変更）。
#
# もともと claude vision だけだった。理由は「作った時点でPDFをOCRする手段がそれしか無かった」
# だけで、必然ではない。実測すると **claude は 186件/2時間**（夜間120分の上限で毎晩186件）で、
# 共有フォルダの残り約6,600件に**35晩以上**かかる。しかも 03:00〜05:00 に定額枠を占有し、
# 2026-08-27 には翻訳ジョブを巻き添えにして全滅させた。
#
# macOS Vision は OS 同梱で**無料・ネットワーク不要・実測 2.3秒/ページ**。
# ただし**手書きや複雑な帳票は苦手**で、表の構造も読めない（文字を拾うだけ）。
# そこで **Vision を先に試し、文字が取れなかったものだけ claude に回す**。
#   ・大半（印刷された書類）は無料・高速で片付く
#   ・手書き等の難物だけ claude が拾う ＝ 枠の消費が桁違いに減る
#
# 実体は mail-archiver で作った道具を借りる（同じMacに1本だけ置く。複製しない）。
# 各PCで `mail-archiver/tools/build.sh` を1回叩けば作れる。無ければ静かに claude へ回る。
VISION_OCR_BIN = os.environ.get(
    "VISION_OCR_BIN",
    os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(
        os.path.abspath(__file__)))), "mail-archiver", "tools", "ocr_pdf"))
# これ未満なら「Vision では読めなかった」とみなして claude へ回す
VISION_MIN_CHARS = 30


def _quota_saver_active() -> bool:
    """claudeの定額枠の節約モード中か（`~/.ai-quota-saver` の1行目が期限）。

    ★**期限つき**にしてある（2026-08-31）。on/off だけだと戻し忘れて永久に止まる。
      日時を過ぎれば勝手に元へ戻る。切り替えは `~/ai-quota-saver.sh`。
    """
    import datetime
    p = os.path.expanduser("~/.ai-quota-saver")
    if not os.path.exists(p):
        return False
    try:
        for line in open(p, encoding="utf-8"):
            line = line.strip()
            if line and not line.startswith("#"):
                # ISO表記なので文字列比較で正しく並ぶ
                return datetime.datetime.now().strftime("%Y-%m-%dT%H:%M") < line
    except OSError:
        return False
    return False


def _ocr_vision(path, max_pages=15) -> list:
    """macOS Vision で文字起こし。取れなければ空リスト（＝claude へ回す合図）。"""
    import subprocess
    if not os.path.exists(VISION_OCR_BIN):
        return []
    try:
        r = subprocess.run([VISION_OCR_BIN, path, "--max-pages", str(max_pages)],
                           capture_output=True, timeout=300)
    except Exception:      # noqa: BLE001 … Vision が転んでも claude で拾えるので止めない
        return []
    if r.returncode != 0:
        return []
    raw = r.stdout.decode("utf-8", "replace")
    out = []
    for i, page in enumerate(raw.split("\x0c"), 1):      # ページ区切りは \f
        body = page.strip()
        if len(body) >= VISION_MIN_CHARS:
            out.append((body, f"P{i}(OCR)"))
    total = sum(len(t) for t, _ in out)
    return out if total >= VISION_MIN_CHARS else []


def ocr_pdf(path, max_pages=15) -> list:
    """スキャン画像PDFを文字起こし。戻り値: [(text, 'P{n}(OCR)'), ...]。

    **二段構え**: ① macOS Vision（無料・高速）→ 取れなければ ② claude vision（手書きに強い）。
    """
    pages = _ocr_vision(path, max_pages=max_pages)
    if pages:
        return pages

    # ★節約モード（2026-08-31）: claude の定額枠が少ないときは回送しない。
    #   ここで OcrUnavailable を投げるのが大事。空リストで返すと呼び出し側が
    #   「この書類には文字が無い」と見なして**見送りリストへ入れ、二度とOCRしなくなる**。
    #   例外なら「今日はできなかった」扱いなので、枠が戻れば自動で再挑戦される。
    if _quota_saver_active():
        raise OcrSkippedByQuotaSaver("節約モード中（Vision では読めなかったので後日に回す）")

    import re
    import tempfile

    from services.claude_client import ClaudeError, run_claude

    with tempfile.TemporaryDirectory() as tmp:
        pages = _render_pdf_images(path, tmp, max_pages=max_pages)
        if not pages:
            return []
        names = "\n".join(f"- page-{n:02d}.png（{n}ページ目）" for n, _ in pages)
        prompt = (
            f"次のスキャン画像ファイル（ディレクトリ {tmp} 内）を Read ツールで1枚ずつ開き、"
            "各ページに書かれている文字を上から順にそのまま文字起こししてください。"
            "解釈・要約・整形は不要で、認識した文字だけを返します。"
            "各ページの冒頭に必ず [P（ページ番号）] を付けてください。\n\n対象:\n" + names
        )
        try:
            env = run_claude(prompt, model="sonnet", timeout=600, add_dir=tmp, allow_read=True)
        except ClaudeError as e:
            # ★空リストで返さない（2026-08-28）。
            #   claude 側の失敗（定額枠切れ・接続不良）を「この書類には文字が無い」と
            #   区別できず、呼び出し側が**見送りリストに登録して永久にスキップ**していた。
            #   2026-08-28 01:00 の夜間OCRで実際に5件がそうなった。
            raise OcrUnavailable(str(e)) from e
        text = (env.get("result") or "").strip()
    if not text:
        return []
    parts = re.split(r"\[P\s*(\d+)\s*\]", text)
    out = []
    if len(parts) > 1:
        for i in range(1, len(parts), 2):
            n = parts[i]
            body = parts[i + 1].strip() if i + 1 < len(parts) else ""
            if body:
                out.append((body, f"P{n}(OCR)"))
    else:
        out = [(text, "OCR")]
    return out


def _chunk(text):
    # NFKC 正規化で半角/全角カナ・英数の揺れを吸収し検索ヒット率を上げる
    text = unicodedata.normalize("NFKC", text.replace("\r", ""))
    chunks, i, n = [], 0, len(text)
    while i < n:
        chunks.append(text[i:i + CHUNK_SIZE])
        i += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def _file_hash(path) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def ingest_file(path, category=None, title=None, force=False, ocr_fallback=False,
                 ocr_max_pages=15, company=None) -> dict:
    """1 ファイルを取込。増分更新: 内容ハッシュが前回と同じならスキップ（変更なし）。

    バージョン管理は filepath 単位（同じ物理ファイルの旧版を無効化し version+1）。
    ocr_fallback=True の場合、テキスト層の無い PDF は claude vision で OCR して取込む。

    company … その文書がどの会社のものか（会社の壁。services/company_scope.py 参照）。
      省略すると DB の既定（大京商事株式会社）になる。
      ★別会社の資料を取り込むときは**必ず指定する**。指定を忘れると大京商事の資料として
        索引され、全体チャットワークから読めてしまう（＝壁を越える）。
    """
    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED_EXT:
        return {"skipped": True, "reason": f"未対応形式 {ext}"}

    mtime = os.path.getmtime(path)
    cur = None
    if not force:
        cur = query_one(
            "SELECT content_hash, source_mtime FROM knowledge_documents WHERE filepath=? AND active=1 "
            "ORDER BY version DESC LIMIT 1",
            (path,),
        )
        # mtime が一致すれば内容ハッシュ計算を省いてスキップ（日次リフレッシュを軽くする）
        if cur and cur["source_mtime"] is not None and abs(cur["source_mtime"] - mtime) < 1:
            return {"skipped": True, "reason": "変更なし(mtime)", "unchanged": True}

    fhash = _file_hash(path)
    # mtime は違っても内容が同じ（保存し直し等）ならスキップ
    if not force and cur and cur["content_hash"] == fhash:
        return {"skipped": True, "reason": "変更なし", "unchanged": True}

    sections = extract(path)
    used_ocr = False
    if not sections and ocr_fallback and ext == ".pdf":
        sections = ocr_pdf(path, max_pages=ocr_max_pages)
        used_ocr = bool(sections)
    if not sections:
        return {"skipped": True, "reason": "テキスト抽出なし（スキャン画像PDF等の可能性）"}

    filename = os.path.basename(path)
    title = title or os.path.splitext(filename)[0]
    category = category or "未分類"
    mime = ext + ("-ocr" if used_ocr else "")

    # 同一 filepath の旧版を無効化し version+1（最新版のみ active=1）
    prev = query_one(
        "SELECT MAX(version) AS v FROM knowledge_documents WHERE filepath=?", (path,)
    )
    version = (prev["v"] + 1) if prev and prev["v"] else 1

    with get_conn() as conn:
        conn.execute("UPDATE knowledge_documents SET active=0 WHERE filepath=?", (path,))
        if company:
            cur = conn.execute(
                "INSERT INTO knowledge_documents (category, title, version, filename, filepath, "
                "mime, content_hash, source_mtime, active, company) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, ?)",
                (category, title, version, filename, path, mime, fhash, mtime, company),
            )
        else:
            cur = conn.execute(
                "INSERT INTO knowledge_documents (category, title, version, filename, filepath, "
                "mime, content_hash, source_mtime, active) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1)",
                (category, title, version, filename, path, mime, fhash, mtime),
            )
        doc_id = cur.lastrowid
        ord_i, total = 0, 0
        for text, ref in sections:
            for ch in _chunk(text):
                if not ch.strip():
                    continue
                ord_i += 1
                src = f"{filename}" + (f" / {ref}" if ref else "")
                conn.execute(
                    "INSERT INTO knowledge_chunks (doc_id, ord, text, source_ref) VALUES (?, ?, ?, ?)",
                    (doc_id, ord_i, ch, src),
                )
                total += 1
    return {"doc_id": doc_id, "chunks": total, "version": version, "skipped": False,
            "mime": mime, "used_ocr": used_ocr}


def deactivate(doc_id: int):
    with get_conn() as conn:
        conn.execute("UPDATE knowledge_documents SET active=0 WHERE id=?", (doc_id,))


def list_documents(active_only=False):
    sql = "SELECT d.*, (SELECT COUNT(*) FROM knowledge_chunks c WHERE c.doc_id=d.id) AS chunk_count " \
          "FROM knowledge_documents d"
    if active_only:
        sql += " WHERE active=1"
    sql += " ORDER BY category, title, version DESC"
    return query(sql)


def stats() -> dict:
    d = query_one("SELECT COUNT(*) AS n FROM knowledge_documents WHERE active=1")
    c = query_one("SELECT COUNT(*) AS n FROM knowledge_chunks")
    return {"documents": d["n"] if d else 0, "chunks": c["n"] if c else 0}


# ---- フォルダ走査・増分リフレッシュ（§「更新分だけ1日1回」） ----
def iter_supported(root, limit=None):
    files = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if not _skip_dir(d)]
        for fn in filenames:
            if fn.startswith(".") or fn.startswith("~$"):
                continue
            if os.path.splitext(fn)[1].lower() in SUPPORTED_EXT:
                files.append(os.path.join(dirpath, fn))
    files.sort()
    return files[:limit] if limit else files


def category_of(root, path):
    rel = os.path.relpath(path, os.path.abspath(root))
    parts = rel.split(os.sep)
    return parts[0] if len(parts) > 1 else "未分類"


def prune_missing(root) -> int:
    """索引済みだが実ファイルが消えた文書を無効化。無効化件数を返す。"""
    root = os.path.abspath(root)
    rows = query(
        "SELECT DISTINCT filepath FROM knowledge_documents WHERE active=1 AND filepath LIKE ?",
        (root + os.sep + "%",),
    )
    n = 0
    with get_conn() as conn:
        for r in rows:
            fp = r["filepath"]
            if fp and not os.path.exists(fp):
                conn.execute("UPDATE knowledge_documents SET active=0 WHERE filepath=?", (fp,))
                n += 1
    return n


def ingest_folder(root, incremental=True, limit=None, progress=None, company=None) -> dict:
    """フォルダを走査し増分取込。incremental=True なら変更ファイルのみ再取込。

    company を渡すと、そのフォルダの文書すべてをその会社のものとして索引する
    （会社の壁。省略時は DB 既定＝大京商事株式会社）。
    ★別会社のフォルダを取り込むときは必ず指定する。
    ★company を指定したときは knowledge_source_root を上書きしない。
      あれは「日次リフレッシュの対象＝大京商事の共有フォルダ」を指す状態なので、
      別会社のフォルダで塗り替えると毎日の取込先が変わってしまう。
    """
    from services.settings import set_state
    root = os.path.abspath(root)
    files = iter_supported(root, limit)
    res = {"total": len(files), "ingested": 0, "unchanged": 0, "skipped": 0,
           "failed": 0, "chunks": 0, "errors": []}
    for i, p in enumerate(files, 1):
        try:
            r = ingest_file(p, category=category_of(root, p), force=not incremental,
                            company=company)
            if r.get("unchanged"):
                res["unchanged"] += 1
            elif r.get("skipped"):
                res["skipped"] += 1
            else:
                res["ingested"] += 1
                res["chunks"] += r.get("chunks", 0)
        except Exception as e:
            res["failed"] += 1
            res["errors"].append(f"{os.path.basename(p)}: {e}")
        if progress:
            progress(i, len(files), p)
    res["pruned"] = prune_missing(root)
    set_state("knowledge_last_refresh", datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    if not company:
        set_state("knowledge_source_root", root)
    return res


def last_refresh() -> str:
    from services.settings import get_state
    return get_state("knowledge_last_refresh", "未実行")
