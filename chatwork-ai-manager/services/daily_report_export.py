"""業務日報の書き出し（Word / Excel）と Chatwork 投稿用の整形。

日報の本文は Markdown（`## 見出し` ＋ `- 箇条書き`）で保存されている。
ここではそれを解釈して、そのまま配れる形（.docx / .xlsx）と、
Chatwork に貼れる形（見出し記法が無いので【】に置換）に変換する。
"""
import datetime
import json
import re

_WD = "月火水木金土日"


def _text_units(text: str) -> int:
    """文字の表示幅を数える（全角=2 / 半角=1）。Excelの列幅と同じ単位。"""
    import unicodedata
    return sum(2 if unicodedata.east_asian_width(c) in ("W", "F", "A") else 1 for c in text)


def _wrapped_lines(text: str, width_units: int) -> int:
    """列幅 width_units のセルに収めたときの行数。改行も数える。"""
    n = 0
    for ln in (text or "").split("\n"):
        n += max(1, -(-_text_units(ln) // max(1, width_units)))
    return n


def parse_sections(body: str):
    """本文Markdownを [(見出し, [行, ...]), ...] に分解する。"""
    sections, heading, lines = [], None, []
    for raw in (body or "").splitlines():
        line = raw.rstrip()
        m = re.match(r"^#{2,4}\s*(.+?)\s*$", line)
        if m:
            if heading is not None:
                sections.append((heading, lines))
            heading, lines = m.group(1), []
            continue
        if not line.strip():
            continue
        lines.append(re.sub(r"^[-*・]\s*", "", line.strip()))
    if heading is not None:
        sections.append((heading, lines))
    return sections


def date_label(date_str: str) -> str:
    d = datetime.date.fromisoformat(date_str)
    return f"{d.year}年{d.month}月{d.day}日（{_WD[d.weekday()]}）"


def stats_of(row) -> dict:
    try:
        return json.loads(row["stats"] or "{}")
    except (TypeError, ValueError):
        return {}


def stats_label(row) -> str:
    s = stats_of(row)
    return (f"Chatwork発言 {s.get('messages_own', 0)}件 ／ 本日動いたTODO {s.get('tasks_moved', 0)}件"
            f"（うち完了 {s.get('tasks_done_today', 0)}件）／ 未完了TODO {s.get('tasks_open', 0)}件")


# ---- Chatwork 投稿用 ---------------------------------------------------------
def chatwork_body(row, date_str: str, include_opinion: bool = False) -> str:
    """Chatwork に貼れる形。Chatwork は見出し記法が無いので `## X` を `【X】` にする。"""
    out = [f"📝 業務日報（{date_label(date_str)}・AI作成のたたき台）", ""]
    for heading, lines in parse_sections(row["body"]):
        if not include_opinion and "AI所見" in heading:
            continue
        out.append(f"【{heading}】")
        out += [f"・{ln}" for ln in lines] or ["・特になし"]
        out.append("")
    return "\n".join(out).strip()


# ---- Word --------------------------------------------------------------------
def _set_font(run, name="游ゴシック", size=None, bold=False):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.font.name = name
    # 日本語は eastAsia を指定しないと別フォントになる（python-docx の既知の作法）
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def _table_borders(table, color="808080", sz=6):
    """罫線をXMLで明示する。"Table Grid" スタイル頼みだとビューアによって線が出ない。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def build_docx(date_str: str, rows, path: str) -> str:
    """1人1ページの Word 日報を書き出す。"""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(1.8)
        s.left_margin = s.right_margin = Cm(2.0)

    for i, row in enumerate(rows):
        p = doc.add_paragraph()
        if i:
            p.paragraph_format.page_break_before = True   # 空段落を挟まない（□が残る）
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run("業 務 日 報"), size=18, bold=True)

        # 日付・氏名の枠
        t = doc.add_table(rows=2, cols=4)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False          # 人によって表の幅が変わらないよう固定する
        _table_borders(t)
        for (col, label, value) in ((0, "日　付", date_label(date_str)),
                                    (2, "氏　名", str(row["person"]))):
            _set_font(t.cell(0, col).paragraphs[0].add_run(label), size=10, bold=True)
            _set_font(t.cell(0, col + 1).paragraphs[0].add_run(value), size=10)
        _set_font(t.cell(1, 0).paragraphs[0].add_run("要　約"), size=10, bold=True)
        c = t.cell(1, 1).merge(t.cell(1, 3))
        _set_font(c.paragraphs[0].add_run(row["summary"] or "-"), size=10)
        for w, col in ((Cm(2.2), 0), (Cm(6.0), 1), (Cm(2.2), 2), (Cm(6.0), 3)):
            for cell in t.columns[col].cells:
                cell.width = w

        doc.add_paragraph()
        for heading, lines in parse_sections(row["body"]):
            hp = doc.add_paragraph()
            hp.paragraph_format.space_before = Pt(10)
            hp.paragraph_format.space_after = Pt(2)
            _set_font(hp.add_run(f"■ {heading}"), size=11, bold=True)
            for ln in (lines or ["特になし"]):
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.space_after = Pt(0)
                bp.paragraph_format.left_indent = Cm(1.1)
                bp.paragraph_format.first_line_indent = Cm(-0.55)
                _set_font(bp.add_run(ln), size=10)

        fp = doc.add_paragraph()
        fp.paragraph_format.space_before = Pt(14)
        r = fp.add_run(stats_label(row))
        _set_font(r, size=8)
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(path)
    return path


# ---- Excel -------------------------------------------------------------------
def build_xlsx(date_str: str, rows, path: str) -> str:
    """1シートに全員分を縦に並べた Excel 日報（渡された rows の順に並べる）。"""
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    thin = Side(style="thin", color="AAAAAA")
    box = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill("solid", fgColor="E8EEF7")
    name_fill = PatternFill("solid", fgColor="D6E1F2")
    title_font = Font(name="游ゴシック", size=16, bold=True)
    name_font = Font(name="游ゴシック", size=12, bold=True)
    label_font = Font(name="游ゴシック", size=10, bold=True)
    body_font = Font(name="游ゴシック", size=10)

    wb = Workbook()
    # ※ 標準スタイル（wb._named_styles["Normal"]）のフォントは書き換えないこと。
    #   openpyxl の内部APIで、触ると Excel がファイルを修復扱いにして
    #   題字「業務日報」などの書式が落ちる（2026-08-21 に実際に起きた）。
    #   列幅の単位ずれは、行の高さを多めに見積もることで吸収する。
    ws = wb.active
    ws.title = f"業務日報{date_str[5:].replace('-', '')}"
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 74

    ws["A1"] = "業務日報"
    ws["A1"].font = title_font
    ws.merge_cells("A1:B1")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28
    ws["A2"] = "日付"
    ws["A2"].font = label_font
    ws["A2"].fill = head_fill
    ws["B2"] = date_label(date_str)
    ws["B2"].font = body_font
    for c in ("A2", "B2"):
        ws[c].border = box
        ws[c].alignment = Alignment(vertical="center")

    r = 4
    for row in rows:
        ws.cell(r, 1, str(row["person"]))
        ws.cell(r, 1).font = name_font
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        for c in (1, 2):
            ws.cell(r, c).fill = name_fill
            ws.cell(r, c).border = box
        ws.cell(r, 1).alignment = Alignment(vertical="center")
        ws.row_dimensions[r].height = 22
        r += 1

        for heading, lines in parse_sections(row["body"]):
            ws.cell(r, 1, heading).font = label_font
            ws.cell(r, 1).fill = head_fill
            text = "\n".join(("特になし" if ln == "特になし" else f"・{ln}")
                             for ln in (lines or ["特になし"]))
            ws.cell(r, 2, text).font = body_font
            for c in (1, 2):
                ws.cell(r, c).border = box
                ws.cell(r, c).alignment = Alignment(vertical="top", wrap_text=True)
            # 折り返し行数から高さを出す。B列の幅は74なので、余白を見て 68 単位で折り返すとみなす。
            # （Excelは高さ未指定でも自動調整することがあるが、当てにすると文字が切れる）
            # 高さは Excel の自動調整の実測に合わせている（2026-08-21 に実機で採寸）:
            #   1行 = 18pt ／ 折り返しは B列の幅と同じ 74 単位（全角=2・半角=1）で起きる。
            # 実測9ブロック中8つが一致、残り1つは1行多く見積もる（＝切れない側）。
            n = _wrapped_lines(text, 74)
            ws.row_dimensions[r].height = max(18, min(600, n * 18))
            r += 1
        r += 1   # 人と人のあいだを1行あける

    ws.sheet_view.showGridLines = False
    ws.print_options.horizontalCentered = True
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToWidth = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    wb.save(path)
    return path
