#!/usr/bin/env python3
"""添付ファイルの**中身**を取り出して全文検索に載せる（2026-08-31 追加）。

    /usr/bin/python3 attach_extract.py --stats            進み具合だけ見る
    /usr/bin/python3 attach_extract.py --max-minutes 45   夜間用（時間で切る）
    /usr/bin/python3 attach_extract.py --no-ocr           テキスト層のあるものだけ
    /usr/bin/python3 attach_extract.py --retry-empty      「文字なし」をもう一度だけ試す

## なぜ要るか

これまで添付は「保管してダウンロードできる」だけで、**中身は検索できなかった**。
2026-08-31 に実際に困った例: PTA大会の会場「スイスホテル南海大阪」は
**メール本文に一度も出てこず**、スキャンPDFの中にしかない。どんな語で検索しても当たらなかった。

添付は 39,726件（PDF 26,669 / doc 4,052 / xls 1,681 / xlsx 1,573 / docx 1,112 / 画像 3,700ほか）。
PDFを60件抜き取って調べたところ **58%はテキスト層あり・42%はスキャン画像**だった。

## 3段構え（2026-09-02 にOCRを二段構えへ。オーナー指示）

1. **テキスト層・Office** … その場で取り出す（速い）
2. **スキャン画像** … macOS Vision で OCR（`tools/ocr_pdf`）。無料・ネット不要・実測2.3秒/ページ
3. **Vision で読めなかったものだけ** claude vision（手書き・崩れた帳票に強い）

★もともと 2 で止めていた（claude を使わなかった）。理由は AI業務マネージャーの夜間OCRと
  **同じ定額枠を取り合う**から。3 を足しても枠への影響が小さいのは、**回るのが
  「Visionが落とした分」だけ**だから。共有フォルダ側の実測では、Vision で 8割以上が片付く。

★claude へ回すときの決まり（共有フォルダ側で事故った教訓をそのまま持ってきた）:
  - **節約モード中は claude を呼ばない**。ただし `none`（文字なし）で記録もしない。
    記録すると「この添付には文字が無い」と確定して**枠が戻っても二度と読まれない**
    （2026-08-28 に AI業務マネージャーで実際に起きた事故と同じ型）。行を書かずに次の晩へ回す
  - **claude 側の失敗（枠切れ・接続不良）も `none` にしない**。同じ理由で行を書かない

## 中断と再開

**1添付1行**を `attachment_texts` に必ず書く（文字が取れなくても `method='none'` で記録）。
書かないと毎晩同じファイルを試して前に進まない（`ocr_ingest.py` の教訓）。

★`--limit` と `--max-new` は別物。`--limit` は対象リストの先頭N件を切り出すだけなので、
  毎晩流すと同じ先頭を舐めて終わる。夜間は `--max-minutes` / `--max-new` を使うこと。

## 実行するPython

`/usr/bin/python3`（pdfplumber / openpyxl / python-docx / xlrd が入っている）。
`.venv` には入っていないので、**夜間ジョブからも /usr/bin/python3 で呼ぶ**
（`translate_english.py` と同じ流儀）。ライブラリが無い形式は黙って飛ばさず「未対応」と記録する。
"""
from __future__ import annotations

import argparse
import concurrent.futures as futures
import datetime
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from typing import Optional, Tuple

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import config  # noqa: E402
import db  # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
OCR_BIN = os.path.join(HERE, "tools", "ocr_pdf")

# 1添付あたりの上限。索引DBが際限なく膨らむのを防ぐ（1.6GB→どこまで増えるかを測れるように）
MAX_CHARS = 200_000
# これ未満ならテキスト層が無いとみなす（表紙だけ文字があるPDFを拾わないため）
MIN_TEXT_CHARS = 30
OCR_MAX_PAGES = 20
OCR_TIMEOUT = 240
# ★2段目（claude vision）。Vision がこの文字数に届かなかったものだけ回す。
CLAUDE_BIN = shutil.which("claude") or "/opt/homebrew/bin/claude"
CLAUDE_OCR_TIMEOUT = 600
CLAUDE_OCR_MAX_PAGES = 10      # 1件が長くなりすぎないように。共有フォルダ側(15)より控えめ
QUOTA_SAVER_FILE = os.path.expanduser("~/.ai-quota-saver")

TEXT_EXTS = [".pdf", ".docx", ".xlsx", ".xlsm", ".xls", ".doc",
             ".csv", ".txt", ".md", ".html", ".htm"]
IMAGE_EXTS = [".jpg", ".jpeg", ".png", ".heic", ".tif", ".tiff"]


def _clean(s: str) -> str:
    """空白と改行を詰める。trigram索引に無駄な空白を入れない。"""
    s = re.sub(r"[ \t　]+", " ", s or "")
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()[:MAX_CHARS]


# ---------------------------------------------------------------- 取り出し

def _read_text_file(path: str) -> str:
    raw = open(path, "rb").read()
    for enc in ("utf-8", "cp932", "euc-jp"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", "replace")


def _extract_pdf(path: str) -> Tuple[str, Optional[int]]:
    import pdfplumber
    import warnings
    warnings.filterwarnings("ignore")
    with pdfplumber.open(path) as pdf:
        pages = pdf.pages[:OCR_MAX_PAGES]
        return "\n".join((p.extract_text() or "") for p in pages), len(pdf.pages)


def _extract_docx(path: str) -> Tuple[str, Optional[int]]:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:                       # 表の中に本文があることが多い（案内状・様式）
        for row in t.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts), None


def _extract_xlsx(path: str) -> Tuple[str, Optional[int]]:
    import openpyxl
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception:
        # ★openpyxl は一部の .xlsx で「could not read stylesheet」で開けない
        #   （実データ約6,000件のうち20件＝0.3%。書式が壊れているだけで中身は読める）。
        #   .xlsx は zip なので、中の XML から文字だけ拾って検索には載せる。
        return _extract_xlsx_raw(path), None
    parts = []
    for ws in wb.worksheets:
        parts.append("# " + str(ws.title))
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) for v in row if v is not None]
            if vals:
                parts.append("\t".join(vals))
    wb.close()
    return "\n".join(parts), None


def _extract_xlsx_raw(path: str) -> str:
    """壊れた .xlsx から、zip の中のXMLを直接読んで文字だけ拾う（予備の手）。"""
    import zipfile
    import xml.etree.ElementTree as ET
    out = []
    with zipfile.ZipFile(path) as z:
        names = [n for n in z.namelist()
                 if n == "xl/sharedStrings.xml" or n.startswith("xl/worksheets/sheet")]
        for n in names:
            try:
                root = ET.fromstring(z.read(n))
            except ET.ParseError:
                continue
            for el in root.iter():
                # <t>（共有文字列）と <v>（値）に中身が入っている
                if el.tag.rsplit("}", 1)[-1] in ("t", "v") and (el.text or "").strip():
                    out.append(el.text.strip())
    return "\n".join(out)


def _extract_xls(path: str) -> Tuple[str, Optional[int]]:
    # ★xlrd 2.x は .xlsx を捨てて .xls 専用になった（新しい方は openpyxl が見る）
    import xlrd
    book = xlrd.open_workbook(path)
    parts = []
    for sh in book.sheets():
        parts.append("# " + sh.name)
        for r in range(sh.nrows):
            vals = [str(v) for v in sh.row_values(r) if str(v).strip()]
            if vals:
                parts.append("\t".join(vals))
    return "\n".join(parts), None


def _extract_doc(path: str) -> Tuple[str, Optional[int]]:
    """旧Word。macOS 同梱の textutil で取り出す（antiword は入っていない）。"""
    r = subprocess.run(["textutil", "-convert", "txt", "-stdout", path],
                       capture_output=True, timeout=120)
    if r.returncode != 0:
        raise RuntimeError((r.stderr or b"").decode("utf-8", "replace")[:200])
    raw = r.stdout
    for enc in ("utf-8", "cp932"):           # textutil は Shift_JIS のまま返すことがある
        try:
            return raw.decode(enc), None
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", "replace"), None


def _extract_html(path: str) -> Tuple[str, Optional[int]]:
    s = _read_text_file(path)
    s = re.sub(r"(?is)<(script|style).*?</\1>", " ", s)
    s = re.sub(r"(?s)<[^>]+>", " ", s)
    return s, None


_EXTRACTORS = {
    ".pdf": _extract_pdf, ".docx": _extract_docx,
    ".xlsx": _extract_xlsx, ".xlsm": _extract_xlsx,
    ".xls": _extract_xls, ".doc": _extract_doc,
    ".html": _extract_html, ".htm": _extract_html,
    ".csv": lambda p: (_read_text_file(p), None),
    ".txt": lambda p: (_read_text_file(p), None),
    ".md": lambda p: (_read_text_file(p), None),
}


def ocr(path: str) -> str:
    """macOS Vision で文字起こし。土台が無ければ RuntimeError（黙って空にしない）。"""
    if not os.path.exists(OCR_BIN):
        raise RuntimeError("OCRの土台が無い: {}（tools/build.sh で作る）".format(OCR_BIN))
    r = subprocess.run([OCR_BIN, path, "--max-pages", str(OCR_MAX_PAGES)],
                       capture_output=True, timeout=OCR_TIMEOUT)
    if r.returncode != 0:
        raise RuntimeError((r.stderr or b"").decode("utf-8", "replace")[:200])
    return r.stdout.decode("utf-8", "replace").replace("\x0c", "\n")


class OcrPostponed(Exception):
    """claude へ回せなかったので**後日に回す**。異常ではない。

    ★これを `none`（文字なし）と混ぜないことが肝。`none` で行を書くと「この添付には
      文字が無い」と確定し、`--retry-empty` を人が明示しない限り**二度と読まれない**。
      例外にしておけば行を書かないので、次の晩に自然と再挑戦される。
    """


def _quota_saver_active() -> bool:
    """claude の定額枠の節約モード中か（`~/.ai-quota-saver` の1行目が期限）。

    ★期限つき。on/off だけだと戻し忘れて永久に止まる。切り替えは `~/ai-quota-saver.sh`。
    """
    if not os.path.exists(QUOTA_SAVER_FILE):
        return False
    try:
        with open(QUOTA_SAVER_FILE, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith("#"):
                    # ISO表記なので文字列比較で正しく並ぶ
                    return datetime.datetime.now().strftime("%Y-%m-%dT%H:%M") < line
    except OSError:
        return False
    return False


def _render_pdf_images(path: str, out_dir: str, max_pages: int, dpi: int = 150) -> list:
    """PDFの各ページをPNGにする。claude は PDF を直接読めないため。"""
    import fitz
    out = []
    with fitz.open(path) as doc:
        for i, page in enumerate(doc, 1):
            if i > max_pages:
                break
            fp = os.path.join(out_dir, "page-{:02d}.png".format(i))
            page.get_pixmap(dpi=dpi).save(fp)
            out.append((i, fp))
    return out


def _run_claude_ocr(prompt: str, work_dir: str) -> str:
    """claude に画像を読ませて文字起こしさせる。失敗は OcrPostponed（`none` にしない）。"""
    cmd = [CLAUDE_BIN, "-p", prompt, "--output-format", "json",
           "--dangerously-skip-permissions", "--model", "sonnet",
           "--add-dir", work_dir]
    try:
        # env は絞らない（CLAUDECODE 等が要る）／stdin=DEVNULL も付けない
        proc = subprocess.run(cmd, capture_output=True, text=True,
                              timeout=CLAUDE_OCR_TIMEOUT, cwd=work_dir)
    except FileNotFoundError as e:
        raise OcrPostponed("claude コマンドが見つからない") from e
    except subprocess.TimeoutExpired as e:
        raise OcrPostponed("claude が{}秒を超えた".format(CLAUDE_OCR_TIMEOUT)) from e
    if proc.returncode != 0:
        raise OcrPostponed("claude が失敗（code {}）: {}".format(
            proc.returncode, (proc.stderr or "")[:150]))
    try:
        outer = json.loads(proc.stdout)
    except json.JSONDecodeError as e:
        raise OcrPostponed("claude の返答をJSONとして読めなかった") from e
    if outer.get("is_error"):
        raise OcrPostponed("claude がエラーを返した")
    return (outer.get("result") or "").strip()


# ★claude は「文字はありません」と**説明文で**返してくることがある（2026-09-02 に実測。
#   `画像に写っている文字はありません(すべてイラスト・写真のみで…)` がそのまま本文として
#   索引に入った）。説明文が検索に載ると邪魔なので、**合図の語だけを返させて捨てる**。
# ★言い回しは実測で決めた（2026-09-02）。「説明を書くな／無ければ合図を返せ」だけを強く言うと、
#   **文字が写っているのに `[文字なし]` を返す**ようになった（バナー画像の
#   `35〜45cmまで指定可能（1cm単位）` を取り落とした）。**拾う側を先に、強く**言うこと。
NO_TEXT_MARK = "[文字なし]"
_RULES = ("小さな文字・帳票の項目名・数字・注記も含め、**読み取れた文字はすべて**そのまま"
          "書き出してください。解釈・要約・整形は不要で、説明や前置きも書かないこと。"
          "文字がまったく写っていない場合に限り " + NO_TEXT_MARK + " とだけ返してください。")


def _strip_no_text(s: str) -> str:
    """claude が「文字なし」と言ってきたら空にする（説明文を索引に入れない）。"""
    t = (s or "").strip()
    if not t or NO_TEXT_MARK in t:
        return ""
    # 合図を無視して地の文で答えてきた場合の保険。短い断り文だけを落とす
    if len(t) < 120 and re.search(r"(文字|テキスト)(は|が)?(写って|含まれて|見当たり)?"
                                  r"(いません|ありません|見当たりません)", t):
        return ""
    # 前置きの1行（「画像内の文字は以下の通りです。」等）を落とす。
    # 索引に入ると検索の邪魔になるうえ、全部の添付に同じ文が並ぶため（2026-09-02 実測）
    lines = t.split("\n")
    if lines and re.match(r"^.{0,40}(以下の通りです|文字起こし(結果)?です|次のとおりです)[。:：]?$",
                          lines[0].strip()):
        t = "\n".join(lines[1:]).strip()
    return t


def ocr_claude(path: str, is_image: bool) -> str:
    """2段目のOCR。Vision で読めなかったものだけがここへ来る。"""
    if _quota_saver_active():
        raise OcrPostponed("節約モード中（Vision では読めなかったので後日に回す）")
    if is_image:
        with tempfile.TemporaryDirectory() as tmp:
            # 元ファイルを触らせない（添付の実体は原本なので読み取り専用で扱う）
            fp = os.path.join(tmp, "page-01" + os.path.splitext(path)[1].lower())
            shutil.copy2(path, fp)
            return _strip_no_text(_run_claude_ocr(
                "{} を Read ツールで開き、写っている文字を上から順にそのまま"
                "文字起こししてください。{}".format(fp, _RULES), tmp))
    with tempfile.TemporaryDirectory() as tmp:
        pages = _render_pdf_images(path, tmp, max_pages=CLAUDE_OCR_MAX_PAGES)
        if not pages:
            return ""
        names = "\n".join("- page-{:02d}.png（{}ページ目）".format(n, n) for n, _ in pages)
        return _strip_no_text(_run_claude_ocr(
            "次のスキャン画像ファイル（ディレクトリ {} 内）を Read ツールで1枚ずつ開き、"
            "各ページに書かれている文字を上から順にそのまま文字起こししてください。{}"
            "\n\n対象:\n{}".format(tmp, _RULES, names), tmp))


def _ocr_two_stage(path: str, is_image: bool) -> Tuple[str, str]:
    """① Vision →（読めなければ）② claude。戻り値 (method, text)。

    method は 'ocr'（Vision で読めた）/ 'ocr-claude'（2段目で読めた）/ 'none'（両方だめ）。
    """
    # ★Vision が転んでも「失敗」で終わらせず claude へ回す（2026-09-02）。
    #   添付の実体は Dropbox(CloudStorage) にあるので、**未ダウンロードのファイルを開くところで
    #   長時間ブロックして時間切れになる**ことがある（共有フォルダ側で実際に一晩を潰した型）。
    #   ここで error にすると行が書かれ、`--retry-empty` を人が明示するまで読まれない。
    try:
        t = _clean(ocr(path))
    except (subprocess.TimeoutExpired, RuntimeError):
        t = ""
    if len(t) >= MIN_TEXT_CHARS:
        return "ocr", t
    t2 = _clean(ocr_claude(path, is_image))     # 失敗時は OcrPostponed が飛ぶ
    # ★2段目は MIN_TEXT_CHARS(30) の下限を当てない（2026-09-02）。
    #   1段目にこの下限があるのは「表紙だけ文字があるPDFをテキスト層ありと誤判定しない」ため。
    #   2段目は事情が違って、**もう枠を使って読んだあと**。短くても中身は本物なので捨てない。
    #   実例: バナー画像から取れた `35〜45cmまで指定可能（1cm単位）` は22文字で、
    #   下限を当てると「文字なし」で確定し、**枠を使ったのに検索にも出ない**という最悪の形になる。
    if t2:
        return "ocr-claude", t2
    # ★ここだけが本当の「文字が無い」。両方の目で見て何も取れなかったので none で確定させる
    return "none", t


def extract_one(path: str, filename: str, use_ocr: bool) -> Tuple[str, str, Optional[int], str]:
    """1件ぶん。戻り値 (method, text, pages, error)。**例外は投げない**。"""
    ext = os.path.splitext(filename)[1].lower()
    if not os.path.exists(path):
        return "error", "", None, "実体が無い"
    try:
        if ext in IMAGE_EXTS:
            if not use_ocr:
                return "skip", "", None, ""
            method, t = _ocr_two_stage(path, is_image=True)
            return method, t, None, ""

        fn = _EXTRACTORS.get(ext)
        if fn is None:
            return "error", "", None, "未対応の形式: {}".format(ext)

        try:
            raw, pages = fn(path)
        except ImportError as e:              # ライブラリ不足を「文字なし」と混同しない
            return "error", "", None, "ライブラリ不足: {}".format(e)
        text = _clean(raw)

        if ext == ".pdf" and len(text) < MIN_TEXT_CHARS:
            # テキスト層が無い＝スキャン画像。OCRへ回す
            if not use_ocr:
                return "skip", "", pages, ""
            method, t = _ocr_two_stage(path, is_image=False)
            return method, t, pages, ""

        return ("text" if len(text) >= MIN_TEXT_CHARS else "none"), text, pages, ""
    except OcrPostponed as e:
        # ★行を書かずに次の晩へ回す（`none` にしない＝二度と読まれなくなるのを防ぐ）
        return "postpone", "", None, str(e)[:150]
    except subprocess.TimeoutExpired:
        return "error", "", None, "時間切れ"
    except Exception as e:                     # noqa: BLE001 … 1件で全体を止めない
        return "error", "", None, "{}: {}".format(type(e).__name__, str(e)[:150])


# ---------------------------------------------------------------- 本体

def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--max-minutes", type=float, default=0, help="この分数で切り上げる（夜間用）")
    ap.add_argument("--max-new", type=int, default=0, help="処理できた件数がNに達したら終了")
    ap.add_argument("--limit", type=int, default=0, help="対象の先頭N件だけ（テスト用）")
    ap.add_argument("--since-days", type=int, default=0,
                    help="この日数より新しいメールの添付だけ（新着を先に片付ける用）")
    ap.add_argument("--workers", type=int, default=4, help="同時に処理する数")
    ap.add_argument("--no-ocr", action="store_true", help="OCRしない（テキスト層のみ）")
    ap.add_argument("--ocr-only", action="store_true", help="OCRが要るものだけ")
    ap.add_argument("--retry-empty", action="store_true", help="none/error をもう一度試す")
    ap.add_argument("--stats", action="store_true", help="進み具合だけ出して終わる")
    args = ap.parse_args()

    conn = db.connect(config.DB_PATH)
    db.init_schema(conn)                       # 新しい表は CREATE IF NOT EXISTS で足りる

    if args.stats:
        s = db.attachment_text_stats(conn)
        print("添付 {:,}件 / 処理済み {:,}件".format(s["添付総数"], s["処理済み"]))
        for k, v in sorted(s["内訳"].items()):
            print("  {:6s} {:>7,}件  {:>12,}文字".format(k, v["件"], v["文字"]))
        return 0

    exts = (IMAGE_EXTS if args.ocr_only else TEXT_EXTS + IMAGE_EXTS)
    todo = db.pending_attachments(conn, exts, limit=args.limit,
                                  retry_empty=args.retry_empty,
                                  since_days=args.since_days)
    use_ocr = not args.no_ocr
    started = time.time()
    print("{} 添付の中身を索引に入れる … 対象 {:,}件{} / OCR {} / 同時 {}".format(
        time.strftime("%Y-%m-%d %H:%M:%S"), len(todo),
        "（直近{}日ぶん）".format(args.since_days) if args.since_days else "",
        "する" if use_ocr else "しない", args.workers), flush=True)

    counts = {"text": 0, "ocr": 0, "ocr-claude": 0, "none": 0,
              "error": 0, "skip": 0, "postpone": 0}
    done = 0
    stop_reason = "対象をすべて処理した"

    def work(row):
        path = os.path.join(config.DATA_DIR, row["path"])
        return row, extract_one(path, row["filename"], use_ocr)

    with futures.ThreadPoolExecutor(max_workers=max(1, args.workers)) as pool:
        it = iter(todo)
        pending = set()
        # ★全部を一度に投げない。--max-minutes で切るとき、投げた分だけ無駄に走るため
        for _ in range(args.workers * 2):
            nxt = next(it, None)
            if nxt is None:
                break
            pending.add(pool.submit(work, nxt))

        while pending:
            finished, pending = futures.wait(pending, return_when=futures.FIRST_COMPLETED)
            for f in finished:
                row, (method, text, pages, err) = f.result()
                # ★postpone は行を書かない。書くと「処理済み」になって次の晩に来ない
                if method not in ("skip", "postpone"):
                    db.save_attachment_text(conn, row["id"], row["message_id"],
                                            method, text, pages, err)
                    if counts.get(method) is not None:
                        done += 1
                counts[method] = counts.get(method, 0) + 1
            # ★毎回コミットする（2026-08-31）。まとめてコミットにすると、
            #   **書き込みトランザクションが開いたまま次の抽出を待つ**ことになる。
            #   OCRが混ざると1回の待ちが数分になり、その間ずっとDBがロックされて
            #   閲覧UI（8535）が "database is locked" で落ちた（実際に発生）。
            #   WAL なのでコミット自体は軽い。ロックを長く持たないことのほうが大事。
            conn.commit()
            if done and done % 50 < len(finished):
                print("  {:,}件 … text {:,} / ocr {:,} / 文字なし {:,} / 失敗 {:,}".format(
                    done, counts["text"], counts["ocr"], counts["none"],
                    counts["error"]), flush=True)

            over_time = args.max_minutes and (time.time() - started) / 60 >= args.max_minutes
            over_new = args.max_new and done >= args.max_new
            if over_time or over_new:
                stop_reason = "時間切れ（{}分）".format(args.max_minutes) if over_time \
                    else "上限（{}件）に達した".format(args.max_new)
                for f in pending:
                    f.cancel()
                break
            for _ in range(len(finished)):
                nxt = next(it, None)
                if nxt is None:
                    break
                pending.add(pool.submit(work, nxt))

    conn.commit()
    mins = (time.time() - started) / 60
    print("{} 終了 {:.1f}分 / {}".format(time.strftime("%Y-%m-%d %H:%M:%S"), mins, stop_reason))
    print("  テキスト層 {:,} / OCR(Vision) {:,} / OCR(claude) {:,} / 文字なし {:,} / "
          "失敗 {:,} / 見送り {:,}".format(
              counts["text"], counts["ocr"], counts["ocr-claude"], counts["none"],
              counts["error"], counts["skip"]))
    if counts["postpone"]:
        print("  後日に回した {:,}件（節約モード中か claude が応答しなかった分）。"
              "★行を書いていないので次の晩に自動で再挑戦する".format(counts["postpone"]))
    s = db.attachment_text_stats(conn)
    print("  進み具合: {:,} / {:,}件（{:.1f}%）".format(
        s["処理済み"], s["添付総数"], 100.0 * s["処理済み"] / max(1, s["添付総数"])))
    return 0


if __name__ == "__main__":
    sys.exit(main())
