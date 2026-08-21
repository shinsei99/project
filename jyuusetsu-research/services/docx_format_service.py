"""Word 形式の公式書式（賃貸借契約書・媒介契約書など）への流し込み。

Excel 側（official_format_service）と役割は同じだが、構造がまるで違うので分けている。

  Excel … 入力欄が「色」で区別され、他書式へ数式で波及する
  Word  … 入力欄は**表のセル**。色も数式も無く、見出しの隣のセルが入力欄

書類雛形フォルダの内訳（2026-08-21 実測）:
  xlsx 74本 / docx 112本 / doc 14本
  → **賃貸借契約書36本と媒介契約書10本は Word しか無い**ので、この経路が要る。
  `.doc`（旧Word・14本）は python-docx で読めない。Word で .docx 保存し直しが要る。

書き込みの作法:
  段落の run を全部消して書き直すと**フォント・下線・網掛けが飛ぶ**ので、
  先頭 run のテキストだけ差し替え、残りの run は空文字にする（書式は先頭 run のものが残る）。
"""

from __future__ import annotations

import os
import re
from typing import Dict, List, Optional

from docx import Document

from services.field_map import normalize

# 見出し（表の左セル）→ PropertyData のキー。
# Excel 側の RULES とは別に持つ。Word の賃貸借契約書は語彙が違うため
# （「名称」「所在地」「構造」「種類」「床面積」など、重説より素朴）。
LABEL_RULES = [
    ("所在地",   r"所在地|所在$",     r"事務所|本店"),
    ("構造",     r"構造",             r"形状"),
    ("種類",     r"種類",             r"権利の種類"),
    ("床面積",   r"床面積|専有面積",   None),
    ("名称",     r"名称|物件名",       r"商号"),
    ("地番",     r"地番",             r"家屋番号"),
    ("地目",     r"地目",             None),
    ("地積",     r"地積",             r"確定|実測"),
    ("家屋番号", r"家屋番号",         None),
    ("所有者",   r"登記名義人|名義人|貸主|賃貸人", None),
]

# 入力欄に既に入っている案内文（これらは消さずに後ろへ足す）
_GUIDE = re.compile(r"^[（(].{1,8}[）)]$")


def _cell_texts(row) -> List[str]:
    return [c.text.strip().replace("\n", "") for c in row.cells]


def _distinct_cells(row):
    """結合セルは python-docx が同じオブジェクトを繰り返し返すので畳む。"""
    out = []
    for c in row.cells:
        if not out or c._tc is not out[-1]._tc:
            out.append(c)
    return out


def scan(path: str) -> List[dict]:
    """表を走査し、{field, table, row, col, current} の候補を返す。"""
    doc = Document(path)
    found: List[dict] = []
    used = {}  # field -> 最初に当たった表の番号
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            cells = _distinct_cells(row)
            texts = [c.text.strip().replace("\n", "") for c in cells]
            for ci, t in enumerate(texts):
                n = normalize(t)
                if not n:
                    continue
                for field, inc, exc in LABEL_RULES:
                    # 同じ表の中なら同じ項目を何度でも当てる。
                    # 賃貸借契約書の建物表示は所在地が「(住居表示)」と「(登記簿)」の
                    # 2行に分かれており、1回で打ち切ると登記簿側が空のまま出る。
                    # 別の表に出てきたものは別項目の可能性が高いので拾わない。
                    if field in used and used[field] != ti:
                        continue
                    if not re.search(inc, n):
                        continue
                    if exc and re.search(exc, n):
                        continue
                    # 見出しの右にある最初の別セルが入力欄
                    for cj in range(ci + 1, len(cells)):
                        found.append({
                            "field": field, "table": ti, "row": ri, "col": cj,
                            "label": t[:24], "current": texts[cj][:30],
                        })
                        used.setdefault(field, ti)
                        break
                    break
    return found


def _write_cell(cell, value: str, keep_guide: bool) -> None:
    """セルへ書き込む。書式を保つため先頭 run だけ差し替える。"""
    paras = cell.paragraphs
    p = paras[0]
    guide = p.text.strip() if keep_guide and _GUIDE.match(p.text.strip()) else ""
    text = ("%s %s" % (guide, value)).strip() if guide else value
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in paras[1:]:
        for r in extra.runs:
            r.text = ""


def fill(src_path: str, dst_path: str, data: Dict[str, str],
         targets: Optional[List[dict]] = None) -> str:
    """PropertyData を Word 書式へ流し込む。

    空の項目は触らない（書式の案内文・選択肢をそのまま残すため）。
    """
    doc = Document(src_path)
    targets = targets if targets is not None else scan(src_path)
    for t in targets:
        value = str(data.get(t["field"], "") or "").strip()
        if not value:
            continue
        cell = _distinct_cells(doc.tables[t["table"]].rows[t["row"]])[t["col"]]
        _write_cell(cell, value, keep_guide=True)
    os.makedirs(os.path.dirname(dst_path), exist_ok=True)
    doc.save(dst_path)
    return dst_path
