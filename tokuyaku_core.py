# -*- coding: utf-8 -*-
"""特約条項の本文生成と書き出し — リポジトリ直下の共有モジュール。

2026-08-21 に `tokuyaku-generator/app.py` から実体をここへ移した。
`tokuyaku-generator`（8513・社内稼働中）と `jyuusetsu-research`（重説）の
**両方がこの1本を読む**。画面（Streamlit）はそれぞれのアプリに残し、
ここには**本文生成・Word組版・テキスト組立てだけ**を置く。

**コピーを作らないこと。** 条文の作り方が2箇所に分かれると、
片方だけ直した特約が契約書に載る。

移設時に `CLAUDE_BIN` の固定パスを直した。`/opt/homebrew/bin/claude` 決め打ちだったため、
Intel Mac や `~/.local/bin` にCLIを入れているPCでは見つからず、
**AI生成が使えなかった**（2026-08-21 サブPCで判明）。
"""

import json
import os
import shutil
import subprocess
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from tokuyaku_clauses import CATEGORIES, all_items, find_item  # noqa: F401


def _find_claude() -> str:
    """`claude` の実体を探す。固定パスにしない（PCごとに場所が違う）。"""
    found = shutil.which("claude")
    if found:
        return found
    for cand in ("~/.local/bin/claude", "/opt/homebrew/bin/claude", "/usr/local/bin/claude"):
        path = os.path.expanduser(cand)
        if os.path.exists(path):
            return path
    return "claude"


CLAUDE_BIN = _find_claude()
CLAUDE_TIMEOUT = 120


STYLE_GUIDE = {
    "である調（契約書標準）": "文末は「〜とする」「〜するものとする」等の常体（である調）で統一する。",
    "ですます調": "文末は丁寧な敬体（ですます調）で統一する。",
}


# ── AI生成（claude CLI） ──────────────────────────────────────────────────────
def generate_clause(item: dict, ctx: dict, style: str, extra: str) -> str:
    """1つの特約項目について本文（条文）を生成する。"""
    prop = ctx.get("property", "").strip()
    seller = ctx.get("seller", "").strip()
    buyer = ctx.get("buyer", "").strip()

    parts = []
    if prop:
        parts.append(f"対象物件: {prop}")
    if seller:
        parts.append(f"売主の表記: {seller}")
    if buyer:
        parts.append(f"買主の表記: {buyer}")
    ctx_block = "\n".join(parts) if parts else "（物件情報の指定なし。一般的な表記で作成）"

    style_rule = STYLE_GUIDE.get(style, STYLE_GUIDE["である調（契約書標準）"])
    extra_block = f"\n■ 追加の事情・条件:\n{extra.strip()}" if extra.strip() else ""

    prompt = f"""あなたは不動産売買契約の特約条項作成に精通したベテラン宅地建物取引士です。
以下の項目について、不動産売買契約書にそのまま挿入できる「特約条項の本文」を作成してください。

■ 特約項目: {item['category']} ＞ {item['title']}
■ 検索キーワード: {item['hint']}
■ 物件情報:
{ctx_block}{extra_block}

【作成ルール】
- 出力は特約条項の本文のみ。見出し番号・タイトル・解説・前置き・後書きは含めない。
- できるだけ簡潔に。原則1項（1〜3文程度）でまとめ、内容上どうしても必要な場合のみ2項までとする。2項にする場合は「1.」「2.」と項番号を付ける。
- 一般的な不動産売買契約の特約文の標準的な粒度・長さに合わせ、冗長な言い回しや同義の繰り返しを避け、要点のみを端的に記載する。
- {style_rule}
- 該当する法令名・条番号があれば正確に引用する（建築基準法第42条第2項 等）。
- 物件情報が指定されていれば自然に織り込み、未指定の箇所は「本物件」「売主」「買主」等の一般表記にする。
- 不明な数値・固有名詞は創作せず、金額・距離・日数等は「〇〇」とプレースホルダにする。

特約条項の本文のみ（簡潔に）を出力してください:"""

    cmd = [
        CLAUDE_BIN, "-p", prompt,
        "--output-format", "json",
        "--dangerously-skip-permissions",
        "--model", "sonnet",
    ]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=CLAUDE_TIMEOUT)
    except FileNotFoundError:
        raise RuntimeError("`claude` コマンドが見つかりません。Claude Code CLI がインストールされているか確認してください。")
    except subprocess.TimeoutExpired:
        raise RuntimeError(f"生成が{CLAUDE_TIMEOUT}秒を超えたため中断しました。再試行してください。")

    if proc.returncode != 0:
        raise RuntimeError(f"claude コマンドが失敗しました（終了コード {proc.returncode}）\n{proc.stderr.strip()[:300]}")

    result = json.loads(proc.stdout)
    if result.get("is_error"):
        raise RuntimeError(f"Claude がエラーを返しました: {result.get('result')}")
    return result.get("result", "").strip()


# ── Word出力 ─────────────────────────────────────────────────────────────────
def _set_font(run, size_pt=10.5, bold=False, color=None, font="游明朝"):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font)
    if color:
        run.font.color.rgb = RGBColor(*color)


def build_docx(clauses: list, ctx: dict) -> bytes:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(title.add_run("特 約 条 項"), size_pt=15, bold=True, font="游ゴシック")

    if ctx.get("property", "").strip():
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(sub.add_run(f"物件：{ctx['property'].strip()}"), size_pt=10, font="游ゴシック")

    doc.add_paragraph()

    for idx, c in enumerate(clauses, 1):
        head = doc.add_paragraph()
        _set_font(head.add_run(f"第{idx}条（{c['title']}）"), size_pt=11, bold=True, font="游ゴシック")
        body_text = (c.get("text") or "（本文未生成）").strip()
        for line in body_text.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            _set_font(p.add_run(line), size_pt=10.5)
        doc.add_paragraph()

    note = doc.add_paragraph()
    _set_font(
        note.add_run("※本書はAIが作成した下書きです。必ず専門家によるリーガルチェックと表記統一を行ってください。"),
        size_pt=8, color=(150, 150, 150),
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def assemble_text(clauses: list) -> str:
    blocks = []
    for idx, c in enumerate(clauses, 1):
        body = (c.get("text") or "（本文未生成）").strip()
        blocks.append(f"第{idx}条（{c['title']}）\n{body}")
    return "\n\n".join(blocks)


