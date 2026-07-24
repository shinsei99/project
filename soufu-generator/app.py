import streamlit as st
import subprocess
import json
import os
import io
import html as html_mod
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CLAUDE_BIN = "/opt/homebrew/bin/claude"
CLAUDE_TIMEOUT = 120
SENDERS_FILE  = os.path.join(os.path.dirname(os.path.abspath(__file__)), "senders.json")

DEFAULT_SENDERS = [
    {
        "label": "大京",
        "company": "大京商事株式会社",
        "title": "専務取締役",
        "name": "鷲見　慎一",
        "address": "大阪市都島区東野田町2-3-14",
        "tel": "０６－６３５３－０４１８",
        "fax": "０６－６３５３－０２８０",
    },
    {
        "label": "新誠",
        "company": "新誠プロパティマネジメント㈱",
        "title": "代表取締役",
        "name": "鷲見　慎一",
        "address": "大阪市北区大淀中3-1-15",
        "tel": "０６－６９３５－７２６７",
        "fax": "０６－７６３５－７８１１",
    },
    {
        "label": "京橋",
        "company": "一般社団法人京橋地域活性化機構",
        "title": "理事長",
        "name": "鷲見　慎一",
        "address": "大阪市都島区東野田町3-10-3",
        "tel": "０６－６９３５－７２６７",
        "fax": "０６－７６３５－７８１１",
    },
    {
        "label": "個人",
        "company": "",
        "title": "",
        "name": "鷲見　慎一",
        "address": "大阪市城東区中央1-10-22-901",
        "tel": "090-8530-0184",
        "fax": "０６－７６３５－７８１１",
    },
]

PREAMBLE = (
    "拝啓　時下いよいよご清祥のこととお慶び申し上げます。\n"
    "平素は何かとご高配を賜り有難く厚くお礼申し上げます。\n"
    "さて、掲題につき下記の通り添付申し上げますので\n"
    "ご査収下さいますようお願い申し上げます。　　　　敬具"
)

_FW  = str.maketrans("0123456789", "０１２３４５６７８９")
# 全角数字・記号 → 半角（TEL/FAX 表示用）
_HALF = str.maketrans("０１２３４５６７８９－（）　", "0123456789-() ")

_ASCII_FONT = "Times New Roman"   # 英数字フォント
_JP_FONT    = "ＭＳ 明朝"           # 日本語フォント（Windows標準の明朝）


def to_halfwidth(s: str) -> str:
    return (s or "").translate(_HALF)


# ── Persistence ───────────────────────────────────────────────────────────────
def load_senders() -> list:
    if os.path.exists(SENDERS_FILE):
        try:
            with open(SENDERS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return DEFAULT_SENDERS[:]


def save_senders(senders: list):
    with open(SENDERS_FILE, "w", encoding="utf-8") as f:
        json.dump(senders, f, ensure_ascii=False, indent=2)


# ── Date ──────────────────────────────────────────────────────────────────────
def to_reiwa(d: date) -> str:
    y = d.year - 2018
    return f"令和{str(y).translate(_FW)}年{str(d.month).translate(_FW)}月{str(d.day).translate(_FW)}日"


# ── AI Generation ─────────────────────────────────────────────────────────────
def generate_body(recipient: str, memo: str, sender: dict) -> str:
    sender_name = f"{sender.get('title', '')}　{sender.get('name', '')}".strip()
    sender_desc = f"{sender['company']} {sender_name}" if sender.get("company") else sender_name

    prompt = f"""ビジネス文書の専門家として、送付状の「記」以下に書く本文メッセージを作成してください。

■ 送付先: {recipient}
■ 送付者: {sender_desc}
■ 送付の概要・用件: {memo.strip() if memo.strip() else "（詳細未入力）"}

【作成ルール】
- 「お世話になります。」で書き出す
- 何を送付するか、理由・背景を自然な流れで1〜3文にまとめる
- 締めは「よろしくお願いいたします。」など状況に合わせて
- 余計な前置きや署名は一切不要
- 送付状らしい簡潔・丁寧な文体
- 本文のみを出力すること

本文のみを出力してください:"""

    cmd = [CLAUDE_BIN, "-p", prompt, "--output-format", "json",
           "--dangerously-skip-permissions", "--model", "sonnet"]
    try:
        proc = subprocess.run(cmd, capture_output=True, text=True, timeout=CLAUDE_TIMEOUT)
    except FileNotFoundError:
        raise RuntimeError("`claude` コマンドが見つかりません。")
    except subprocess.TimeoutExpired:
        raise RuntimeError(f"生成が{CLAUDE_TIMEOUT}秒を超えたため中断しました。")
    if proc.returncode != 0:
        raise RuntimeError(f"claude コマンドが失敗しました\n{proc.stderr.strip()[:300]}")
    result = json.loads(proc.stdout)
    if result.get("is_error"):
        raise RuntimeError(f"Claude がエラーを返しました: {result.get('result')}")
    return result.get("result", "").strip()


# ── Word Document（python-docx でゼロから生成 = 常に正当な OOXML）──────────────
# 旧実装は変換ツール由来の非標準テンプレート(template.docx)を差し込む方式で、
# 不正な要素/属性名により Word が「破損 → 開いて修復」を要求していた。
# python-docx で新規生成すれば styles.xml 等を含む正当なパッケージになり破損しない。
# レイアウトは完成イメージプレビュー（render_preview）に合わせている。

def _sender_lines(sender: dict) -> list:
    """プレビューと同一の送付者行 [(テキスト, 太字), ...]（TEL/FAXは半角）。"""
    lines = []
    if sender.get("company"):
        lines.append((sender["company"], True))
    tn = "　".join(filter(None, [sender.get("title", ""), sender.get("name", "")]))
    if tn:
        lines.append((tn, False))
    if sender.get("address"):
        lines.append((sender["address"], False))
    if sender.get("tel"):
        lines.append((f"TEL  {to_halfwidth(sender['tel'])}", False))
    if sender.get("fax"):
        lines.append((f"FAX  {to_halfwidth(sender['fax'])}", False))
    return lines


def _style_run(run, size_pt: float, bold: bool = False):
    run.font.size = Pt(size_pt)
    run.font.name = _ASCII_FONT
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), _JP_FONT)


def _add_para(doc, text: str, size_pt: float, align, bold: bool = False,
              space_after: float = 4):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = 1.15
    _style_run(p.add_run(text), size_pt, bold)
    return p


def _add_rule_borders(p):
    """段落の上下に罫線（「書類送付の件」の囲み枠を再現）。"""
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "bottom"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "4")
        e.set(qn("w:color"), "555555")
        pbdr.append(e)
    pPr.append(pbdr)


def create_docx(recipient: str, doc_date: date, body: str, sender: dict) -> io.BytesIO:
    """完成イメージプレビュー準拠の送付状を python-docx で新規生成して返す。"""
    R = WD_ALIGN_PARAGRAPH.RIGHT
    L = WD_ALIGN_PARAGRAPH.LEFT
    C = WD_ALIGN_PARAGRAPH.CENTER

    doc = Document()

    # ページ設定（Letter 8.5×11 / 余白 上下1" 左右1.25"）
    sec = doc.sections[0]
    sec.page_width   = Inches(8.5)
    sec.page_height  = Inches(11)
    sec.top_margin   = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin  = Inches(1.25)
    sec.right_margin = Inches(1.25)

    # 既定フォント（Normal スタイル）にも日本語フォントを設定
    normal = doc.styles["Normal"]
    normal.font.name = _ASCII_FONT
    normal.font.size = Pt(12)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), _JP_FONT)

    _add_para(doc, to_reiwa(doc_date), 12, R)                       # 日付（右）
    _add_para(doc, f"{recipient}　様", 16, L, bold=True, space_after=10)  # 宛名（左・大）

    for text, bold in _sender_lines(sender):                        # 送付者（右寄せ）
        _add_para(doc, text, 12, R, bold=bold, space_after=2)

    _add_para(doc, "", 12, L, space_after=8)                        # 余白

    title = _add_para(doc, "書 類　送　付 の 件", 14, C, bold=True, space_after=12)
    _add_rule_borders(title)                                        # 上下罫線

    for line in PREAMBLE.split("\n"):                               # 前文（中央）
        _add_para(doc, line, 12, C, space_after=2)

    _add_para(doc, "記", 13, C, bold=True, space_after=8)           # 記（中央）

    for line in [l for l in body.split("\n") if l.strip()]:         # 本文（左）
        _add_para(doc, line, 12, L, space_after=4)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# ── Preview HTML ──────────────────────────────────────────────────────────────
def render_preview(recipient: str, doc_date: date, body: str, sender: dict):
    e = html_mod.escape

    s_lines = []
    if sender.get("company"):
        s_lines.append(f"<b>{e(sender['company'])}</b>")
    tn = "　".join(filter(None, [sender.get("title", ""), sender.get("name", "")]))
    if tn:
        s_lines.append(e(tn))
    if sender.get("address"):
        s_lines.append(e(sender["address"]))
    if sender.get("tel"):
        s_lines.append(f"TEL&nbsp;&nbsp;{e(to_halfwidth(sender['tel']))}")
    if sender.get("fax"):
        s_lines.append(f"FAX&nbsp;&nbsp;{e(to_halfwidth(sender['fax']))}")

    st.markdown(
        f"""
<div style="
    background:white;border:1px solid #ccc;border-radius:6px;
    font-family:'Hiragino Mincho Pro','MS 明朝','Yu Mincho',serif;
    color:#111;box-shadow:0 3px 12px rgba(0,0,0,0.1);
    max-width:680px;margin:0 auto;padding:28px 36px 24px 36px;
    line-height:2;font-size:13px;">
  <div style="text-align:right;margin-bottom:10px;">{e(to_reiwa(doc_date))}</div>
  <div style="margin-bottom:6px;font-size:17px;font-weight:bold;">{e(recipient)}　様</div>
  <div style="margin-bottom:16px;line-height:1.9;text-align:right;">{"<br>".join(s_lines)}</div>
  <div style="text-align:center;font-size:15px;font-weight:bold;letter-spacing:3px;
      border-top:1px solid #555;border-bottom:1px solid #555;
      padding:6px 0;margin-bottom:14px;">書 類　送　付 の 件</div>
  <div style="margin-bottom:14px;text-align:center;">{e(PREAMBLE).replace(chr(10),"<br>")}</div>
  <div style="text-align:center;font-size:14px;font-weight:bold;margin:14px 0 10px;">記</div>
  <div style="margin-bottom:20px;">{e(body).replace(chr(10),"<br>")}</div>
</div>
""",
        unsafe_allow_html=True,
    )


# ── Sidebar: Sender Management ─────────────────────────────────────────────────
def sidebar_senders() -> dict:
    senders = st.session_state.setdefault("senders", load_senders())

    st.markdown("## 📮 送付者を選択")
    labels = [s["label"] for s in senders]
    idx = st.session_state.get("sender_idx", 0)
    if idx >= len(senders):
        idx = 0

    chosen_label = st.radio("送付者", labels, index=idx, key="sender_radio",
                            label_visibility="collapsed")
    chosen_idx = labels.index(chosen_label)
    st.session_state["sender_idx"] = chosen_idx
    sender = senders[chosen_idx]

    st.divider()
    with st.expander("✏️ 送付者を編集・追加", expanded=False):
        mode = st.radio("操作", ["選択中を編集", "新規追加"], horizontal=True, key="edit_mode")
        et = senders[chosen_idx].copy() if mode == "選択中を編集" else \
             {"label": "", "company": "", "title": "", "name": "", "address": "", "tel": "", "fax": ""}

        e_label   = st.text_input("ラベル",   value=et["label"],   key="e_label")
        e_company = st.text_input("会社名",   value=et["company"], key="e_company")
        e_title   = st.text_input("役職",     value=et["title"],   key="e_title")
        e_name    = st.text_input("氏名",     value=et["name"],    key="e_name")
        e_address = st.text_input("住所",     value=et["address"], key="e_address")
        e_tel     = st.text_input("ＴＥＬ",  value=et["tel"],     key="e_tel")
        e_fax     = st.text_input("ＦＡＸ",  value=et["fax"],     key="e_fax")

        new_data = {"label": e_label, "company": e_company, "title": e_title,
                    "name": e_name, "address": e_address, "tel": e_tel, "fax": e_fax}

        col1, col2 = st.columns(2)
        with col1:
            if st.button("💾 保存", use_container_width=True):
                if not e_label.strip():
                    st.warning("ラベルを入力してください")
                else:
                    if mode == "新規追加":
                        senders.append(new_data)
                        st.session_state["sender_idx"] = len(senders) - 1
                    else:
                        senders[chosen_idx] = new_data
                    save_senders(senders)
                    st.session_state["senders"] = senders
                    st.success("保存しました")
                    st.rerun()
        with col2:
            if mode == "選択中を編集" and len(senders) > 1:
                if st.button("🗑 削除", use_container_width=True, type="secondary"):
                    senders.pop(chosen_idx)
                    st.session_state["sender_idx"] = 0
                    save_senders(senders)
                    st.session_state["senders"] = senders
                    st.rerun()

    return sender


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    st.set_page_config(page_title="送付書ジェネレーター", page_icon="📮", layout="wide")

    with st.sidebar:
        sender = sidebar_senders()

    st.title("📮 送付書ジェネレーター")
    st.caption("宛名と用件を入力するだけで、AIが送付状を自動生成 → Word（.docx）ダウンロード")
    st.divider()

    left, right = st.columns([1, 1.2], gap="large")

    with left:
        st.subheader("📝 送付内容を入力")
        recipient = st.text_input("① 宛名", placeholder="例：山田　太郎",
                                  help="「様」は自動で付きます。")
        doc_date  = st.date_input("② 日付", value=date.today())
        memo = st.text_area(
            "③ 送付内容・用件メモ",
            placeholder="例：\n・契約書類一式をお送りします\n・先日お話しした見積書を添付します",
            height=150,
            help="箇条書きでもOK。AIが自然な文章に仕上げます。",
        )
        gen_btn = st.button("✨ 送付書を生成する", type="primary", use_container_width=True)

    with right:
        st.subheader("📄 プレビュー・ダウンロード")

        if gen_btn:
            if not recipient.strip():
                st.warning("⚠️ 宛名を入力してください。")
            elif not memo.strip():
                st.warning("⚠️ 送付内容・用件メモを入力してください。")
            else:
                with st.spinner("AIが本文を生成しています..."):
                    try:
                        body = generate_body(recipient.strip(), memo, sender)
                        st.session_state["body"]        = body
                        st.session_state["meta"]        = {"recipient": recipient.strip(),
                                                           "doc_date": doc_date}
                        st.session_state["has_content"] = True
                        st.success("✅ 生成完了！")
                    except Exception as ex:
                        st.error(f"❌ エラーが発生しました: {ex}")

        if st.session_state.get("has_content"):
            meta   = st.session_state["meta"]
            edited = st.text_area("✏️ 本文（記 以下）を確認・編集", height=160, key="body",
                                  help="自由に編集できます。")

            st.markdown("---")
            with st.expander("📋 完成イメージ（プレビュー）", expanded=True):
                render_preview(meta["recipient"], meta["doc_date"], edited, sender)

            st.markdown("---")
            try:
                buf = create_docx(meta["recipient"], meta["doc_date"], edited, sender)
                from datetime import datetime
                fname = f"送付書_{meta['recipient']}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.download_button(
                    label="📥 Word（.docx）をダウンロード",
                    data=buf,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary",
                )
                st.caption(fname)
            except Exception as ex:
                st.error(f"Word生成エラー: {ex}")

        elif not gen_btn:
            st.info("👈 左側で宛名・用件を入力し、「送付書を生成する」ボタンを押してください。")


if __name__ == "__main__":
    main()
