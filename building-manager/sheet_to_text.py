#!/usr/bin/env python3
"""Office系ファイル（xlsx/xls/docx）を claude が読めるテキストに変換する。

claude CLI の Read ツールはバイナリの Office ファイルを直接読めないため、
AI抽出の前処理として本スクリプトでプレーンテキスト(.txt)へ変換する。

使い方: python3 sheet_to_text.py <入力ファイル> <出力txtパス>
成功で 0、非対応/失敗で非0を返す（呼び出し側は非0なら元ファイルのまま続行）。
"""
import sys
import warnings

warnings.filterwarnings("ignore")


def xlsx_to_text(path: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"=== シート: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip() != ""]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def xls_to_text(path: str) -> str:
    import xlrd
    wb = xlrd.open_workbook(path)
    out = []
    for ws in wb.sheets():
        out.append(f"=== シート: {ws.name} ===")
        for r in range(ws.nrows):
            cells = [str(ws.cell_value(r, c)).strip() for c in range(ws.ncols)]
            cells = [c for c in cells if c != ""]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def docx_to_text(path: str) -> str:
    import docx
    doc = docx.Document(path)
    out = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def main() -> int:
    if len(sys.argv) < 3:
        print("usage: sheet_to_text.py <input> <output.txt>", file=sys.stderr)
        return 2
    src, dst = sys.argv[1], sys.argv[2]
    low = src.lower()
    try:
        if low.endswith(".xlsx"):
            text = xlsx_to_text(src)
        elif low.endswith(".xls"):
            text = xls_to_text(src)
        elif low.endswith(".docx"):
            text = docx_to_text(src)
        else:
            print(f"unsupported: {src}", file=sys.stderr)
            return 3
    except Exception as e:  # noqa: BLE001
        print(f"convert failed: {e}", file=sys.stderr)
        return 4

    if not text.strip():
        print("empty result", file=sys.stderr)
        return 5
    with open(dst, "w", encoding="utf-8") as f:
        f.write(text)
    return 0


if __name__ == "__main__":
    sys.exit(main())
