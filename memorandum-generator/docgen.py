# -*- coding: utf-8 -*-
"""
覚書・合意書・各種申請書ジェネレーター 文書生成エンジン
大京商事株式会社「覚書・同居申請等」フォルダの170書類を分析しパターン化したもの。

全書類は共通骨格を持つ:
  タイトル → 物件表示 → 当事者(甲/乙/丙)と原契約日 → 記 → 本文条項
  → 「本覚書に定めのない事項は原契約による」 → 作成通数 → 日付欄 → 署名欄(甲乙丙+連帯保証人+立会人)
"""
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

FONT = "ＭＳ 明朝"

# ── 大京商事（立会人・自社情報）デフォルト ─────────────────────────
DAIKYO = {
    "company": "大　京　商　事　株　式　会　社",
    "company_plain": "大京商事株式会社",
    "title": "代表取締役",
    "rep": "鷲　見　　文　子",
    "address": "大阪市都島区東野田町２丁目３番１４号",
    "tel": "０６－６３５３－０４１８",
}


# ── 低レベルヘルパ ──────────────────────────────────────────────
def _font(run, size=10.5, bold=False, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def _p(doc, text="", *, size=10.5, bold=False, align="left", space_after=4, space_before=0, indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    if indent is not None:
        pf.left_indent = Cm(indent)
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    if text:
        for i, line in enumerate(str(text).split("\n")):
            if i > 0:
                run = p.add_run()
                run.add_break()
                _font(run, size, bold)
            run = p.add_run(line)
            _font(run, size, bold)
    else:
        _font(p.add_run(""), size, bold)
    return p


def _era_line(era="令和"):
    if era in ("令和", "平成", "昭和"):
        return f"{era}　　年　　月　　日"
    return "　　年　　月　　日"


# ── 署名欄レンダリング ──────────────────────────────────────────
def _signatures(doc, signatories, witness=True, witness_rep=None):
    """signatories: list of dict(role, name, address, seal=True)"""
    for s in signatories:
        _p(doc, space_after=2)
        _p(doc, s["role"], space_after=2)
        _p(doc, f"住所　　{s.get('address','')}", indent=0.7, space_after=2)
        name = s.get("name", "")
        seal = "　　　　　　　　㊞" if s.get("seal", True) else ""
        _p(doc, f"氏名　　{name}{seal}", indent=0.7, space_after=2)
    if witness:
        rep = witness_rep or DAIKYO["rep"]
        _p(doc, space_after=2)
        _p(doc, "（立会人）", space_after=2)
        _p(doc, f"住所　　{DAIKYO['address']}", indent=0.7, space_after=2)
        _p(doc, f"氏名　　{DAIKYO['company']}", indent=0.7, space_after=2)
        _p(doc, f"　　　　{DAIKYO['title']}　{rep}　　　㊞", indent=0.7, space_after=2)


def _property_block(doc, prop):
    """prop: dict(address, name, area)"""
    _p(doc, f"（物件表示）所在地：{prop.get('address','')}", space_after=2)
    _p(doc, f"　　　　　　名　称：{prop.get('name','')}", space_after=2)
    if prop.get("area"):
        _p(doc, f"　　　　　　契約面積：{prop['area']}", space_after=2)


def _new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return doc


def _save(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── 覚書ファミリー共通フレーム ─────────────────────────────────────
def _memo_frame(title, prop, preamble, body_paragraphs, signatories,
                era="令和", witness=True, witness_rep=None, kishou="・・・・・記・・・・・"):
    doc = _new_doc()
    _p(doc, title, size=16, bold=True, align="center", space_before=6, space_after=14)
    _property_block(doc, prop)
    _p(doc, space_after=2)
    _p(doc, preamble, space_after=8)
    if kishou:
        _p(doc, kishou, align="center", space_after=8)
    for bp in body_paragraphs:
        _p(doc, bp.get("text", ""), indent=bp.get("indent"), space_after=bp.get("space_after", 6))
    _p(doc, space_after=6)
    _p(doc, _era_line(era), space_after=8)
    _signatures(doc, signatories, witness=witness, witness_rep=witness_rep)
    return _save(doc)


# ═══════════════════════════════════════════════════════════════════
# 各書類タイプのビルダー
# 全て inputs(dict) を受け取り .docx バイト列を返す
# ═══════════════════════════════════════════════════════════════════

def _lessor_role(landlord_is_person):
    # 賃貸人/賃借人 vs 貸主/借主  … サンプル上どちらも使われるが賃貸人/賃借人を既定
    return ("賃貸人", "賃借人")


def build_rent_revision(d):
    """賃料改定（現行/改定・賃料/共益費/水道代）"""
    lr, lt = "賃貸人", "賃借人"
    preamble = (
        f"{lr}　{d['ko_name']}（以下甲という）と{lt}　{d['otsu_name']}（以下乙という）の両者は、"
        f"{d['orig_date']}付賃貸借期間開始の頭書物件の賃貸借契約（以下原契約という）について、"
        "原契約の賃料及び共益費・水道代を次の通り変更する事に合意した。"
    )
    def money_block(label, rent, kyoueki, suido):
        lines = [label, f"賃　料／月額　　{rent}"]
        if kyoueki:
            lines.append(f"共益費／月額　　{kyoueki}")
        if suido:
            lines.append(f"水道代／月額　　{suido}")
        if d.get("show_total") and (kyoueki or suido):
            lines.append(f"　　　　　　　　計{d.get('total_' + label,'')}")
        return "\n".join(lines)
    body = [
        {"text": "１．現行賃料及び共益費・水道代\n" +
                 f"賃　料／月額　　{d['cur_rent']}" +
                 (f"\n共益費／月額　　{d['cur_kyoueki']}" if d.get('cur_kyoueki') else "") +
                 (f"\n水道代／月額　　{d['cur_suido']}" if d.get('cur_suido') else "") +
                 (f"\n　　　　　　　　計{d['cur_total']}" if d.get('cur_total') else "")},
        {"text": "２．改定賃料及び共益費・水道代\n" +
                 f"賃　料／月額　　{d['new_rent']}" +
                 (f"\n共益費／月額　　{d['new_kyoueki']}" if d.get('new_kyoueki') else "") +
                 (f"\n水道代／月額　　{d['new_suido']}" if d.get('new_suido') else "") +
                 (f"\n　　　　　　　　計{d['new_total']}" if d.get('new_total') else "")},
        {"text": f"３．本覚書の開始日は{d['start_date']}とし、本覚書に定めのない事項については、"
                 "原契約に定めるところとする。"},
        {"text": "上記合意を証するため本書を２通作成し、甲乙各１通づつ保有する。\n以上", "space_after": 2},
    ]
    sig = [
        {"role": "賃貸人（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "賃借人（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
    ]
    return _memo_frame("覚　　書", d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", False),
                       witness_rep=d.get("witness_rep"))


def build_rent_reduction(d):
    """家賃値下げ・賃料減額（シンプル版）"""
    preamble = (
        f"貸主　{d['ko_name']}（以下甲という）と借主　{d['otsu_name']}（以下乙という）の両者は、"
        f"{d['orig_date']}締結の頭書物件の賃貸借契約（以下原契約という）について、"
        "原契約の賃料を次の通り変更する事に合意した。"
    )
    kback = d.get("kyoueki_note", "共益費は賃料に込み")
    body = [
        {"text": f"１．賃料及び共益費（{kback}）"},
        {"text": f"{d['start_date']}からは下記のとおりとし、以降の改定については原契約のとおりとする。"},
        {"text": f"賃　料／月額　　{d['new_rent']}"},
        {"text": "将来において、消費税率に変更が生じた場合は、適用日を基準に変更するものとする。"},
        {"text": "２．本覚書に定めのない事項については、原契約に定めるところとする。"},
        {"text": "上記合意を証するため本書を２通作成し、甲乙各１通を保有する。\n以上", "space_after": 2},
    ]
    sig = [
        {"role": "貸主（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "借主（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
    ]
    return _memo_frame("覚　　書", d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", True),
                       witness_rep=d.get("witness_rep"))


def build_succession(d):
    """契約上の地位承継（甲乙丙＋新連帯保証人）
    前文で甲/乙/丙を定義し、本文条項は「甲」「乙」「丙」の略称のみで記載する。"""
    ko, otsu, hei = d["ko_name"], d["otsu_name"], d["hei_name"]
    # 新連帯保証人は既定で旧賃借人(乙)。別人を指定した場合のみその表記を用いる。
    hosho = d.get("hosho_name") or "乙"
    preamble = (
        f"賃貸人　{ko}（以下「甲」という）と賃借人　{otsu}（以下「乙」という）及び{hei}（以下「丙」という）は、"
        f"甲・乙間において締結した、{d['orig_date']}付の頭書物件の賃貸借契約"
        "（以下原契約という）の契約上の地位承継に関して以下のとおり覚書を締結する。"
    )
    body = [
        {"text": f"（地位承継の期日）\n　　乙は{d['succ_date']}（以下承継期日という）をもって原契約の"
                 "賃借人としての地位を丙に承継させるものとし、甲はこれを承諾する。"},
        {"text": "（契約当事者の読み替え）\n　　甲・丙は、承継期日以降、原契約上の賃借人乙を"
                 "丙と読み替え、原契約を甲・丙間に適用する。"},
        {"text": "（契約書の引継ぎ）\n　　丙は、原契約の原本を乙より引き継いだ。"},
        {"text": "（権利義務の承継）\n　　丙は承継期日をもって、原契約に基づき乙が甲に対して有する、"
                 "権利または負担する一切の義務を乙より承継する。"},
        {"text": "（確認条項）\n　　原契約上の地位の承継に関し乙・丙間に紛争が発生した場合は、"
                 "乙及び丙の全責任と費用負担においてこれを処理解決することとする。"},
        {"text": f"（連帯保証人）\n　　地位継承に伴い{hosho}は新連帯保証人となり、原契約に基づき丙が将来"
                 "甲に対して負担する一切の債務につき丙と連帯して保証し、債務の履行の責に任ずる。"
                 "甲はこれを承認する。"},
        {"text": "　本覚書各条項以外は、原契約の通りとする。"},
        {"text": "本覚書締結の証として本書を４通作成し、甲・乙・新連帯保証人・丙各々記名捺印のうえ"
                 "各自１通を保有する。\n以上", "space_after": 2},
    ]
    sig = [
        {"role": "賃貸人（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "賃借人（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
        {"role": "新賃借人（丙）", "name": d.get("hei_sign", ""), "address": d.get("hei_addr", "")},
        {"role": "新連帯保証人", "name": d.get("hosho_sign", ""), "address": d.get("hosho_addr", "")},
    ]
    return _memo_frame("契約上の地位承継に関する覚書", d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", True),
                       witness_rep=d.get("witness_rep"), kishou="")


def build_rep_change(d):
    """代表取締役変更に伴う連帯保証（甲乙丙）"""
    ko, otsu = d["ko_name"], d["otsu_name"]
    preamble = (
        f"貸主　{ko}（以下甲という）と借主　{otsu}（以下乙という）及び新代表取締役　{d.get('hei_name','')}"
        f"（以下「丙」という）は、丙が乙の代表取締役に就任することに伴い、{d['orig_date']}締結の頭書物件の"
        "賃貸借契約（これに付帯する覚書を含め以下原契約という）について、下記事項について合意を確認し、"
        "本日覚書を締結する。"
    )
    body = [
        {"text": f"１．丙は、{d.get('assume_date','　　年　　月　　日')}をもって原契約の連帯保証人としての地位を承諾し、"
                 "甲及び乙はこれを追認する。"},
        {"text": "２．本覚書に定めのない事項については、原契約に定めるところとする。"},
        {"text": "上記合意を証するため本書を３通作成し、甲・乙・丙各々記名押印のうえ各１通を保有する。\n以上",
         "space_after": 2},
    ]
    sig = [
        {"role": "貸主（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "借主（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
        {"role": "新連帯保証人（丙）", "name": d.get("hei_sign", ""), "address": d.get("hei_addr", "")},
    ]
    return _memo_frame("覚　　書", d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", True),
                       witness_rep=d.get("witness_rep"))


def build_guarantor_delete(d):
    """連帯保証人削除"""
    ko, otsu = d["ko_name"], d["otsu_name"]
    preamble = (
        f"{ko}（以下「甲」という）と　{otsu}　（以下「乙」という）は、{d['orig_date']}付にて締結した"
        "賃貸借契約（以下「原契約」という）に関し、今般、甲乙合意のうえ下記のとおり覚書"
        "（以下「本覚書」という）を交換する。"
    )
    clause_no = d.get("guarantor_clause", "第２４条")
    body = [
        {"text": f"原契約{clause_no}に定める連帯保証人に関する条項を削除することを甲乙双方確認するものとする。"},
        {"text": f"連帯保証人条項を削除するに伴い、「{d.get('guarantor_person','')}」氏の連帯保証人としての地位を"
                 "解除することを甲乙双方確認する。"},
        {"text": "本覚書に定めのない事項についてはすべて「原契約」の定めのとおりとする。"},
        {"text": "上記の特約を証するために本覚書を２通作成し、甲乙記名押印のうえ、各々その１通を保有する。\n以上",
         "space_after": 2},
    ]
    sig = [
        {"role": "（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
    ]
    return _memo_frame("覚　　書", d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", True),
                       witness_rep=d.get("witness_rep"))


def build_restoration(d):
    """原状回復義務の免除（改装項目）"""
    ko, otsu = d["ko_name"], d["otsu_name"]
    preamble = (
        f"賃貸人　{ko}（以下甲という）と賃借人　{otsu}（以下乙という）との間において、"
        f"{d['orig_date']}付締結した頭書物件の賃貸借契約（以下原契約）に関して、下記事項について確認し"
        "合意をしたので、ここに覚書を作成し各自１通ずつ所持するものとする。"
    )
    items = [x.strip() for x in d.get("items", "").split("\n") if x.strip()]
    clause = d.get("resto_clause", "第１９条「貸室の原状回復と明渡し」")
    body = [
        {"text": f"１．原契約書　{clause}について"},
        {"text": f"{d.get('reform_when','')}に甲の承認を得て乙の費用負担にて改装をする下記の項目について、"
                 "乙は原状回復の義務を負わないものとする。尚、乙の費用負担にて新設、改装・変更したものについては"
                 "甲に代価を求めてはならない。"},
    ]
    for it in items:
        body.append({"text": f"　・{it}", "space_after": 2})
    body += [
        {"text": f"その他の事項は原契約書　{clause}のとおりとする。"},
        {"text": "２．本覚書に定めのない事項についてはすべて「原契約」の定めのとおりとする。"},
    ]
    sig = [
        {"role": "賃貸人（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "賃借人（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
    ]
    return _memo_frame("覚　　書", d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", True),
                       witness_rep=d.get("witness_rep"), kishou="記")


def build_parking_change(d):
    """駐車場位置変更"""
    ko, otsu = d["ko_name"], d["otsu_name"]
    preamble = (
        f"貸主　{ko}と借主　{otsu}の両者は、{d['orig_date']}締結の頭書物件の賃貸借契約"
        "（以下原契約という）について、下記の合意をみたので後日の為覚書を２通作成し各一通を所持するものとする。"
    )
    body = [
        {"text": f"駐車場位置を　{d.get('from_no','')}より　{d.get('to_no','')}へ変更とする。"},
        {"text": f"本覚書の開始は{d['start_date']}からとする。"},
        {"text": "本覚書に定めのない事項については、原契約に定めるところとする。"},
        {"text": "上記合意を証するため本書を２通作成し、甲乙各１通づつ保有する。\n以上", "space_after": 2},
    ]
    sig = [
        {"role": "貸主（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "借主（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
    ]
    return _memo_frame("覚　　書", d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", False),
                       witness_rep=d.get("witness_rep"))


def build_name_change(d):
    """名義変更（新設法人へ）"""
    ko, otsu = d["ko_name"], d["otsu_name"]
    preamble = (
        f"賃貸人　{ko}（以下甲という）と賃借人　{otsu}（以下乙という）の間において、"
        f"{d['orig_date']}付締結した、頭書物件の賃貸借契約（以下原契約という）に関し、次のとおり合意した。"
    )
    body = [
        {"text": f"原契約の{d.get('based_on','特約事項第２項')}のとおり、乙名義を{d.get('reason','新設会社法人登記完了')}"
                 "に伴い新賃借人（以下丙という）に名義変更を行うものとする。"},
        {"text": "本覚書に定めのない事項については、原契約に定めるところとする。"},
        {"text": "上記合意を証する為に本書を３部作成し記名捺印の上、各自１通を保有する。\n以上", "space_after": 2},
    ]
    sig = [
        {"role": "賃貸人（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "旧賃借人（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
        {"role": "新賃借人（丙）", "name": d.get("hei_sign", ""), "address": d.get("hei_addr", "")},
    ]
    return _memo_frame("覚　　書", d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", True),
                       witness_rep=d.get("witness_rep"))


def build_freeform(d):
    """汎用 覚書／合意書（本文を自由入力）"""
    ko, otsu = d["ko_name"], d["otsu_name"]
    doctitle = d.get("doc_title", "覚　　書")
    word = "合意書" if "合意" in doctitle else "覚書"
    preamble = (
        f"賃貸人　{ko}（以下甲という）と賃借人　{otsu}（以下乙という）の間において、"
        f"{d['orig_date']}付締結した、頭書物件の賃貸借契約（以下原契約という）に関し、次のとおり合意した。"
    )
    clauses = [x.strip() for x in d.get("clauses", "").split("\n") if x.strip()]
    body = []
    for i, c in enumerate(clauses, 1):
        body.append({"text": f"{i}．{c}"})
    body.append({"text": f"{len(clauses)+1}．本{word}に定めのない事項については、原契約に定めるところとする。"})
    copies = d.get("copies", "２")
    body.append({"text": f"上記合意を証するため本書を{copies}通作成し、各自１通を保有する。\n以上", "space_after": 2})
    sig = [
        {"role": "（甲）", "name": d.get("ko_sign", ""), "address": d.get("ko_addr", "")},
        {"role": "（乙）", "name": d.get("otsu_sign", ""), "address": d.get("otsu_addr", "")},
    ]
    return _memo_frame(doctitle, d, preamble, body, sig,
                       era=d.get("era", "令和"), witness=d.get("witness", True),
                       witness_rep=d.get("witness_rep"))


# ── 特殊フォーマット（覚書骨格でないもの） ──────────────────────────
def build_cohabitation(d):
    """同居申請書＋同居承諾書"""
    doc = _new_doc()
    _p(doc, "同居申請書", size=15, bold=True, align="center", space_after=10)
    _p(doc, f"{d.get('lessor_to','')}　御中　　　　　　　　　　申請日：　{_era_line(d.get('era','令和'))}",
       space_after=8)
    body = (
        f"{d.get('contract_date','')}付けで締結した「{d.get('property_name','')}」貸室賃貸契約書"
        "（以下賃貸契約書という）に基づき、貴社より賃借している店舗・事務所内に、次の者を"
        f"{d.get('cohab_start','　　年　　月　　日')}より同居させたいと存じます。つきましては、当該賃貸借契約"
        f"第{d.get('article','１１')}条により同居は禁止されておりますが、特別に承諾賜りたくお願い申し上げます。"
        "同居希望者履歴事項全部証明書添付のうえ、申請致します。尚、同居の御承諾を得ました上は、下記事項を"
        "厳正に守り、貴社に対して一切のご迷惑をおかけ致しません。"
    )
    _p(doc, body, space_after=8)
    conditions = [
        "当該賃貸借契約が解除となった場合は、理由の如何に拘わらず賃借人の責任において同居者を退室せしめ、物件の完全なる明渡しをしなければならない。",
        "同居者においては、賃貸人に対し又は賃貸契約書上、もしくは借家法上の権利が何ら発生しないものとする。従って賃貸人に対し何らの権利主張もなし得ない。",
        "同居者は賃借人の関係者であり、関連関係が消滅した場合において同居者は、速やかに借室から退去しなければならない。",
        "同居者の行為により生じた賃貸人に対する損害は、全て賃借人の負担とし速やかに賠償の責に任ずるものとする。",
        "賃借人は、同居者に当該物件の賃借権を譲渡してはならない。",
        "同居者が当該賃貸借契約及び管理規定に違反する行為のあった場合は、全て賃借人の契約違反として処理されても賃借人は一切の異議を申し立てることができないものとする。",
    ]
    for c in conditions:
        _p(doc, f"・{c}", indent=0.5, space_after=4)
    _p(doc, space_after=2)
    _p(doc, f"物件表示　　{d.get('prop_addr','')}", space_after=2)
    _p(doc, f"　　　　　　{d.get('property_name','')}", space_after=6)
    _p(doc, "申請人")
    _p(doc, f"（賃借人）　住所　　{d.get('tenant_addr','')}", space_after=2)
    _p(doc, f"　　　　　　氏名　　{d.get('tenant_name','')}　　　　㊞", space_after=6)
    for i in range(1, 4):
        addr = d.get(f"cohab{i}_addr", "")
        name = d.get(f"cohab{i}_name", "")
        if name or addr:
            _p(doc, f"（同居人）　住所　　{addr}", space_after=2)
            _p(doc, f"　　　　　　氏名　　{name}　　　　㊞", space_after=6)
    # 同居承諾書
    doc.add_page_break()
    _p(doc, "同居承諾書", size=15, bold=True, align="center", space_after=10)
    _p(doc, f"{d.get('tenant_name','')}　御中　　　　　　　　　　　{_era_line(d.get('era','令和'))}", space_after=10)
    _p(doc, f"{_era_line(d.get('era','令和'))}付で貴社より提出された同居承諾願いについて記載内容を条件として承諾致します。",
       space_after=12)
    _p(doc, "賃貸人　　住所", space_after=6)
    _p(doc, "　　　　　　氏名　　　　　　　　　　　　　　　　　　㊞")
    return _save(doc)


def build_minor_consent(d):
    """未成年者同意書"""
    doc = _new_doc()
    _p(doc, "同　意　書", size=15, bold=True, align="center", space_after=12)
    _p(doc, _era_line(d.get("era", "令和")), align="right", space_after=10)
    _p(doc, f"現住所　　{d.get('minor_addr','')}", space_after=4)
    _p(doc, f"未成年者氏名　　{d.get('minor_name','')}", space_after=10)
    _p(doc, "上記未成年者が下記条件の下に、賃貸借契約を締結することについて親権者として同意致します。",
       space_after=10)
    _p(doc, f"物件名　　　{d.get('property_name','')}", space_after=2)
    _p(doc, f"物件住所　　{d.get('prop_addr','')}", space_after=2)
    _p(doc, "賃貸借条件", space_after=2)
    for line in [x for x in d.get("conditions", "").split("\n") if x.strip()]:
        _p(doc, f"　　{line.strip()}", space_after=2)
    _p(doc, space_after=8)
    _p(doc, "親権者", space_after=4)
    _p(doc, f"住所　　{d.get('parent_addr','')}", indent=0.5, space_after=4)
    _p(doc, f"氏名　　{d.get('parent_name','')}　　　　　　㊞", indent=0.5)
    return _save(doc)


def build_use_permit(d):
    """使用許可承諾書"""
    doc = _new_doc()
    _p(doc, "使用許可承諾書", size=15, bold=True, align="center", space_after=12)
    _p(doc, "当社が貴社より賃借している下記事業所の全部を下記目的で下記の者に使用させたいので、承諾願います。",
       space_after=8)
    _p(doc, f"（賃借人）{d.get('tenant_name','')}", space_after=2)
    if d.get("tenant_rep"):
        _p(doc, f"　　　　　{d['tenant_rep']}", space_after=8)
    _p(doc, f"　【賃借物件の表示】　{d.get('prop_addr','')}　{d.get('property_name','')}", space_after=2)
    _p(doc, f"　【使用目的・名称】　{d.get('purpose','')}", space_after=2)
    _p(doc, f"　【使用を行う者】　住所　　{d.get('user_addr','')}", space_after=2)
    _p(doc, f"　　　　　　　　　　氏名　　{d.get('user_name','')}", space_after=10)
    _p(doc, "上記について承諾いたします。", space_after=10)
    _p(doc, _era_line(d.get("era", "令和")), space_after=8)
    _p(doc, f"（賃貸人）　住所　　{d.get('lessor_addr','')}", space_after=2)
    _p(doc, f"　　　　　　氏名　　{d.get('lessor_name','')}　　　　　　㊞")
    return _save(doc)


# ── 書類タイプ登録 ─────────────────────────────────────────────────
DOC_TYPES = {
    "rent_revision":   ("賃料改定（現行→改定・共益費/水道代）", build_rent_revision),
    "rent_reduction":  ("家賃値下げ・賃料減額（シンプル）", build_rent_reduction),
    "succession":      ("契約上の地位承継（甲乙丙＋新連帯保証人）", build_succession),
    "rep_change":      ("代表取締役変更に伴う連帯保証", build_rep_change),
    "guarantor_delete":("連帯保証人削除", build_guarantor_delete),
    "restoration":     ("原状回復義務の免除（改装項目）", build_restoration),
    "parking_change":  ("駐車場位置変更", build_parking_change),
    "name_change":     ("名義変更（新設法人へ）", build_name_change),
    "freeform":        ("汎用 覚書／合意書（本文自由入力）", build_freeform),
    "cohabitation":    ("同居申請書＋同居承諾書", build_cohabitation),
    "minor_consent":   ("未成年者同意書", build_minor_consent),
    "use_permit":      ("使用許可承諾書", build_use_permit),
}
